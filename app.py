#!/usr/bin/env python3
"""
Web app para el Agente Licitador — IMAGINE COMUNICACIÓN ANDALUZA S.L.U
"""

import json
import threading
from datetime import datetime, date
from pathlib import Path
from flask import Flask, jsonify, render_template_string, request
import psycopg2
import psycopg2.extras

BASE = Path(__file__).parent
AGENT_FILE = BASE / "agente_licitador.py"

app = Flask(__name__)
_agent_running = False
_agent_last_result = {"ok": None, "stdout": "", "stderr": "", "ts": None}


def cargar_config() -> dict:
    import os
    cfg_file = BASE / "config.json"
    cfg = json.loads(cfg_file.read_text(encoding="utf-8")) if cfg_file.exists() else {}
    # Variables de entorno tienen prioridad (Vercel)
    if os.environ.get("DATABASE_URL"):
        cfg["database_url"] = os.environ["DATABASE_URL"]
    if os.environ.get("GROK_API_KEY"):
        cfg["grok_api_key"] = os.environ["GROK_API_KEY"]
    return cfg

def get_conn():
    import os
    # Preferir parámetros individuales para evitar problemas con caracteres especiales
    db_url = os.environ.get("DATABASE_URL") or cargar_config().get("database_url", "")
    return psycopg2.connect(db_url, sslmode="require")

@app.get("/api/debug")
def api_debug():
    import os
    db_url = os.environ.get("DATABASE_URL", "NO_ENV_VAR")
    masked = db_url[:30] + "..." if len(db_url) > 30 else db_url
    try:
        conn = get_conn()
        conn.close()
        return jsonify({"ok": True, "db_url_preview": masked})
    except Exception as e:
        return jsonify({"ok": False, "db_url_preview": masked, "error": str(e)})

def load_results() -> dict:
    try:
        conn = get_conn()
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT * FROM licitaciones ORDER BY plazo ASC NULLS LAST")
            rows = cur.fetchall()
        conn.close()
        lics = []
        for r in rows:
            lics.append({
                "_id":    r["id"],
                "titulo": r["titulo"] or "",
                "organo": r["organo"] or "",
                "cpvs":   r["cpvs"] if isinstance(r["cpvs"], list) else json.loads(r["cpvs"] or "[]"),
                "importe": str(r["importe"]) if r["importe"] else "",
                "plazo":   str(r["plazo"]) if r["plazo"] else "",
                "enlace":  r["enlace"] or "",
                "fecha_deteccion": str(r["fecha_deteccion"]) if r["fecha_deteccion"] else "",
            })
        return {"licitaciones": lics, "actualizado": str(datetime.now())}
    except Exception as e:
        print(f"[DB ERROR] {e}")
        return {"licitaciones": [], "actualizado": None}


IRRELEVANT_KEYWORDS = [
    "agencia de viajes",
    "señalización de sendas", "sendas peatonales", "sendas ciclistas",
    "transporte permanente de señales",
    "azafatas/os en diferentes edificios", "azafatas en edificios",
    "certamen ganadero", "certámenes ganaderos",
    "coworking",
    "monitores deportivos",
    "festejos taurinos", "taurino",
    "contratación de orquestas", "grupos musicales para la caseta",
    "catalogación y digitalización de documentación",
    "ti y consultoría",
    "concesión administrativa demanial",
    "feria gastronómica de verano",
    "suministro temporal, instalación, puesta en servicio, control técnico",
    "arrendamiento, sin opción de compra, de equipos",
    "salidas gastronómicas",
    "peregrinos y visitantes",
    "rocódromo",
    "atención, información y asistencia",
]

def es_relevante_titulo(titulo: str) -> bool:
    t = titulo.lower()
    return not any(kw.lower() in t for kw in IRRELEVANT_KEYWORDS)

def es_activa(l: dict) -> bool:
    from datetime import date
    plazo = l.get("plazo", "N/D")
    if not plazo or plazo == "N/D":
        return False
    return plazo >= str(date.today())


# ── API ───────────────────────────────────────────────────────────────────────

@app.get("/api/licitaciones")
def api_licitaciones():
    data = load_results()
    lics = data.get("licitaciones", [])

    solo_activas    = request.args.get("activas", "1") == "1"
    solo_relevantes = request.args.get("relevantes", "1") == "1"

    if solo_activas:
        lics = [l for l in lics if es_activa(l)]
    if solo_relevantes:
        lics = [l for l in lics if es_relevante_titulo(l.get("titulo", ""))]

    q      = request.args.get("q", "").lower()
    cpv    = request.args.get("cpv", "")
    minimo = request.args.get("min", 0, type=float)

    if q:
        lics = [l for l in lics if q in l.get("titulo","").lower() or q in l.get("organo","").lower()]
    if cpv:
        lics = [l for l in lics if any(c.startswith(cpv) for c in l.get("cpvs", []))]
    if minimo:
        def importe_num(l):
            try: return float(l.get("importe") or 0)
            except: return 0
        lics = [l for l in lics if importe_num(l) >= minimo]

    lics.sort(key=lambda l: l.get("plazo") or "9999")

    return jsonify({
        "total": len(lics),
        "actualizado": data.get("actualizado"),
        "licitaciones": lics,
    })


@app.get("/api/stats")
def api_stats():
    data = load_results()
    lics = data.get("licitaciones", [])
    lics = [l for l in lics if es_activa(l) and es_relevante_titulo(l.get("titulo", ""))]
    from collections import Counter

    cpv_counter: Counter = Counter()
    for l in lics:
        for c in set(l.get("cpvs", [])):
            cpv_counter[c[:8]] += 1

    importes = []
    for l in lics:
        try:
            v = float(l.get("importe") or 0)
            if v > 0:
                importes.append(v)
        except:
            pass

    return jsonify({
        "total": len(lics),
        "top_cpvs": cpv_counter.most_common(8),
        "importe_total": sum(importes),
        "importe_medio": round(sum(importes)/len(importes), 2) if importes else 0,
        "importe_max": max(importes) if importes else 0,
        "actualizado": data.get("actualizado"),
    })


def _run_agent():
    global _agent_running, _agent_last_result
    import io, contextlib
    try:
        import agente_licitador
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            agente_licitador.main()
        _agent_last_result = {
            "ok": True,
            "stdout": buf.getvalue()[-4000:],
            "ts": datetime.now().strftime("%H:%M:%S"),
        }
    except Exception as e:
        _agent_last_result = {"ok": False, "msg": str(e), "ts": datetime.now().strftime("%H:%M:%S")}
    finally:
        _agent_running = False

@app.post("/api/ejecutar")
def api_ejecutar():
    global _agent_running
    if _agent_running:
        return jsonify({"ok": False, "msg": "El agente ya está en ejecución"})
    _agent_running = True
    _agent_last_result["ok"] = None  # reset
    threading.Thread(target=_run_agent, daemon=True).start()
    return jsonify({"ok": True, "msg": "Agente iniciado en background"})


CPV_NAMES_PY = {
    "79340000":"Publicidad y marketing","79341000":"Publicidad","79342000":"Marketing",
    "79341200":"Gestión de publicidad","79341400":"Campañas de publicidad","79342200":"Promoción",
    "79341100":"Colocación de publicidad","92210000":"Servicios de radio","92220000":"Servicios de televisión",
    "79822500":"Diseño gráfico","74812240":"Producción vídeos publicidad","92111250":"Producción vídeos información",
    "79800000":"Impresión","72413000":"Diseño web","72415000":"Alojamiento web",
    "79413000":"Consultoría marketing","79416000":"Relaciones públicas","79416200":"Asesoramiento RRPP",
    "79952000":"Organización de eventos","79956000":"Ferias y exposiciones",
}

@app.post("/api/analizar")
def api_analizar():
    import urllib.request as urlreq
    data = request.get_json()
    titulo  = data.get("titulo", "Sin título")
    organo  = data.get("organo", "Desconocido")
    importe = data.get("importe", "N/D")
    plazo   = data.get("plazo", "N/D")
    cpvs    = data.get("cpvs", [])
    enlace  = data.get("enlace", "")

    config  = cargar_config()
    api_key = config.get("grok_api_key", "").strip()
    if not api_key:
        return jsonify({"ok": False, "msg": "Falta la API key de Grok. Añádela en config.json → grok_api_key"})

    cpv_desc = ", ".join(CPV_NAMES_PY.get(c[:8], c) for c in dict.fromkeys(c[:8] for c in cpvs) if c)

    try:
        imp_fmt = f"{float(importe):,.0f} €" if importe and importe not in ("N/D","") else "No especificado"
    except:
        imp_fmt = importe or "No especificado"

    prompt = f"""Eres un experto en contratación pública española. Analiza en detalle esta licitación para IMAGINE COMUNICACIÓN ANDALUZA S.L.U, agencia especializada en publicidad, marketing, diseño gráfico, eventos, relaciones públicas y comunicación digital.

LICITACIÓN:
• Título: {titulo}
• Órgano contratante: {organo}
• Importe estimado: {imp_fmt}
• Plazo de presentación: {plazo}
• Categorías CPV: {cpv_desc or "No especificadas"}
• Enlace: {enlace or "No disponible"}

Proporciona un análisis estructurado con estas secciones exactas (usa los emojis como cabecera de cada bloque):

🎯 RESUMEN EJECUTIVO
Describe qué se licita, para qué organismo y cuál es el objetivo principal del contrato. 3-4 frases concretas.

🛠️ SERVICIOS Y ENTREGABLES REQUERIDOS
Lista detallada de todo lo que hay que realizar o entregar: tipos de piezas, canales, formatos, volúmenes si se mencionan, duración del contrato, posibles prórrogas.

📋 REQUISITOS PARA PRESENTARSE
Lista todos los requisitos obligatorios para poder optar al contrato:
- Requisitos de solvencia económica (facturación mínima, seguros, etc.)
- Requisitos de solvencia técnica (experiencia previa, equipo humano, certificaciones)
- Documentación obligatoria (garantías, registros, habilitaciones profesionales)
- Restricciones o condicionantes específicos

⚡ CRITERIOS DE VALORACIÓN
Explica cómo se puntuará la oferta: criterios objetivos (precio) y subjetivos (calidad, proyecto técnico, etc.) con sus pesos si están disponibles. Qué aspectos son clave para ganar.

💡 ENCAJE CON IMAGINE
Puntuación del 1 al 10 sobre el encaje con el perfil de IMAGINE. Justifica qué capacidades de la agencia son un punto fuerte y qué posibles debilidades o gaps habría que cubrir.

✅ RECOMENDACIÓN
Presentarse / No presentarse / Valorar con cautela. Motivo en 2-3 frases directas.

🚀 PRÓXIMOS PASOS
Lista de 4-6 acciones concretas y ordenadas para preparar la oferta si se decide concurrir (descargar pliego, visita técnica, subcontratas necesarias, etc.).

Responde en español, de forma directa y práctica. Sin introducciones ni conclusiones genéricas. Sé específico y útil."""

    payload = json.dumps({
        "model": "llama-3.3-70b-versatile",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2000,
        "temperature": 0.2,
    }).encode("utf-8")

    import http.client, ssl
    try:
        ctx = ssl.create_default_context()
        conn = http.client.HTTPSConnection("api.groq.com", context=ctx, timeout=45)
        conn.request(
            "POST",
            "/openai/v1/chat/completions",
            body=payload,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
        )
        resp = conn.getresponse()
        body = resp.read().decode("utf-8")
        conn.close()
        if resp.status != 200:
            return jsonify({"ok": False, "msg": f"Groq HTTP {resp.status}: {body[:200]}"})
        result   = json.loads(body)
        analysis = result["choices"][0]["message"]["content"]
        return jsonify({"ok": True, "analysis": analysis})
    except Exception as e:
        return jsonify({"ok": False, "msg": f"Error Groq: {e}"})


@app.post("/api/descargar-analisis")
def api_descargar_analisis():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, re
    from flask import send_file

    data     = request.get_json()
    titulo   = data.get("titulo", "Licitación")
    organo   = data.get("organo", "")
    importe  = data.get("importe", "")
    plazo    = data.get("plazo", "")
    analysis = data.get("analysis", "")

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # Cabecera
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr.add_run("CENTINELA — Análisis de Licitación")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    doc.add_paragraph()

    # Ficha
    ficha = doc.add_paragraph()
    ficha.add_run("📌 Título: ").bold = True
    ficha.add_run(titulo)
    if organo:
        p = doc.add_paragraph()
        p.add_run("🏛️ Órgano: ").bold = True
        p.add_run(organo)
    if importe:
        p = doc.add_paragraph()
        p.add_run("💶 Importe: ").bold = True
        try: p.add_run(f"{float(importe):,.0f} €")
        except: p.add_run(importe)
    if plazo:
        p = doc.add_paragraph()
        p.add_run("📅 Plazo: ").bold = True
        p.add_run(plazo)

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()

    # Parsear secciones del análisis
    SECTION_EMOJIS = re.compile(r'^(🎯|🛠️|📋|⚡|💡|✅|🚀)\s')
    for line in analysis.split('\n'):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if SECTION_EMOJIS.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**',''))
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        elif re.match(r'^[-•*]\s', line) or re.match(r'^\d+\.', line):
            text = re.sub(r'^[-•*\d\.]\s*', '', line).replace('**','')
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            text = line.replace('**','')
            doc.add_paragraph(text)

    # Pie
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = pie.add_run(f"Generado por Centinela · IMAGINE COMUNICACIÓN ANDALUZA")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = re.sub(r'[^\w\s-]', '', titulo)[:50].strip().replace(' ','_')
    return send_file(buf, as_attachment=True,
                     download_name=f"Centinela_{safe}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/api/estado")
def api_estado():
    cache_ids = 0
    try:
        conn = get_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) FROM cache_ids")
            cache_ids = cur.fetchone()[0]
        conn.close()
    except:
        pass
    return jsonify({
        "agente_corriendo": _agent_running,
        "licitaciones_en_cache": cache_ids,
        "ultimo_resultado": _agent_last_result,
    })


# ── Frontend ──────────────────────────────────────────────────────────────────

HTML = r"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Centinela — IMAGINE</title>
<style>
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Helvetica,Arial,sans-serif;background:#fff;min-height:100vh;display:flex;overflow:hidden;}

/* App shell */
.app{display:flex;width:100%;height:100vh;background:#fff;overflow:hidden;}

/* Sidebar */
.sidebar{width:220px;background:#fff;border-right:1px solid #EBEBEB;display:flex;flex-direction:column;flex-shrink:0;height:100vh;overflow:hidden;}
.sidebar-head{padding:16px 16px 10px;display:flex;align-items:center;border-bottom:1px solid #F0F0F0;}
.sidebar-logo{font-size:14px;font-weight:600;color:#111;display:flex;align-items:center;gap:9px;}
.sidebar-logo-icon{width:26px;height:26px;display:flex;align-items:center;justify-content:center;}
.sb-section{padding:10px 10px 4px;overflow-y:auto;flex:1;min-height:0;}
.sb-label{font-size:10px;font-weight:600;color:#AAAAAA;letter-spacing:.8px;text-transform:uppercase;padding:0 6px;margin-bottom:4px;}
.sb-item{display:flex;align-items:center;gap:8px;padding:7px 8px;border-radius:7px;cursor:pointer;font-size:13px;color:#444;transition:background .12s,color .12s;user-select:none;}
.sb-item:hover{background:#F5F5F5;color:#111;}
.sb-item.active{background:#F0EFFD;color:#4A4AE8;font-weight:500;}
.sb-item .sb-icon{font-size:15px;width:18px;text-align:center;flex-shrink:0;}
.sb-item .sb-badge{margin-left:auto;background:#F0F0F0;color:#777;font-size:10px;font-weight:600;padding:1px 6px;border-radius:10px;}
.sb-item.active .sb-badge{background:#DDD8FC;color:#4A4AE8;}
.sb-item .sb-dot{width:7px;height:7px;border-radius:50%;flex-shrink:0;}
.sb-item .sb-dot.red{background:#F87171;}
.sb-item .sb-dot.amber{background:#FBBF24;}
.sb-item .sb-dot.green{background:#34D399;}
.sb-item .sb-dot.blue{background:#60A5FA;}
.sb-footer{flex-shrink:0;padding:12px;border-top:1px solid #F0F0F0;}
.sb-user{display:flex;align-items:center;gap:8px;}
.sb-user-av{width:28px;height:28px;border-radius:50%;background:#4A4AE8;color:#fff;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:600;flex-shrink:0;}
.sb-user-name{font-size:12px;font-weight:500;color:#222;}
.sb-user-email{font-size:10px;color:#999;}

/* Main */
.main{flex:1;display:flex;flex-direction:column;overflow:hidden;background:#FAFAFA;}

/* Topbar */
.topbar{background:#fff;border-bottom:1px solid #EBEBEB;padding:0 24px;height:52px;display:flex;align-items:center;gap:12px;flex-shrink:0;}
.topbar-search{display:flex;align-items:center;gap:6px;background:#F5F5F5;border:1px solid #EBEBEB;border-radius:8px;padding:0 10px;height:32px;flex:1;max-width:300px;}
.topbar-search input{border:none;background:transparent;font-size:13px;color:#333;outline:none;width:100%;}
.topbar-search input::placeholder{color:#BBBBBB;}
.topbar-search .search-icon{color:#BBBBBB;font-size:13px;}
.topbar-sep{flex:1;}
.topbar-actions{display:flex;align-items:center;gap:8px;}
.btn-primary{background:#111;color:#fff;border:none;border-radius:8px;padding:0 16px;height:36px;font-size:13px;font-weight:500;cursor:pointer;display:flex;align-items:center;gap:6px;transition:background .15s;white-space:nowrap;flex-shrink:0;}
.btn-primary:hover{background:#333;}
.btn-primary:disabled{background:#BBB;cursor:not-allowed;}
.btn-primary .spinner{display:none;width:13px;height:13px;border:2px solid rgba(255,255,255,.3);border-top-color:#fff;border-radius:50%;animation:spin .7s linear infinite;}
.btn-primary.loading .spinner{display:block;}
.btn-primary.loading .btn-label{display:none;}
@keyframes spin{to{transform:rotate(360deg);}}

/* Page header */
.page-head{padding:20px 24px 0;}
.page-title-row{display:flex;align-items:center;gap:12px;margin-bottom:20px;}
.page-title{font-size:22px;font-weight:600;color:#111;}
.page-badge{background:#F0EFFD;color:#4A4AE8;font-size:12px;font-weight:500;padding:3px 10px;border-radius:20px;border:1px solid #DDD8FC;}

/* Stats row */
.stats-row{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px;padding:0 24px 20px;}
.stat-card{background:#fff;border:1px solid #EBEBEB;border-radius:10px;padding:14px 16px;}
.stat-card .s-label{font-size:11px;color:#AAAAAA;text-transform:uppercase;letter-spacing:.6px;margin-bottom:6px;display:flex;align-items:center;gap:5px;}
.stat-card .s-val{font-size:16px;font-weight:600;color:#111;line-height:1.2;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}
.stat-card .s-sub{font-size:11px;color:#BBBBBB;margin-top:3px;}
.stat-card.blue .s-val{color:#4A4AE8;}
.stat-card.green .s-val{color:#059669;}
.stat-card.amber .s-val{color:#D97706;}
.stat-card.red .s-val{color:#DC2626;}

/* Toolbar */
.toolbar{padding:0 24px 14px;display:flex;align-items:center;gap:10px;}
.tabs{display:flex;gap:2px;background:#F0F0F0;border-radius:8px;padding:3px;overflow-x:auto;flex-shrink:0;}
.tab{padding:5px 14px;border-radius:6px;font-size:12px;font-weight:500;color:#666;cursor:pointer;transition:background .12s,color .12s;border:none;background:transparent;white-space:nowrap;}
.tab.active{background:#fff;color:#111;box-shadow:0 1px 3px rgba(0,0,0,.1);}
.tb-sep{flex:1;}
.tb-select{height:32px;border:1px solid #EBEBEB;border-radius:7px;background:#fff;padding:0 10px;font-size:12px;color:#444;outline:none;cursor:pointer;}
.filter-btn{height:32px;border:1px solid #EBEBEB;border-radius:7px;background:#fff;padding:0 12px;font-size:12px;color:#444;cursor:pointer;display:flex;align-items:center;gap:5px;}
.filter-btn:hover{background:#F5F5F5;}
.count-pill{background:#F0F0F0;color:#666;font-size:11px;font-weight:600;padding:2px 8px;border-radius:10px;}

/* Card grid */
.cards-area{flex:1;overflow-y:auto;padding:0 24px 24px;display:flex;flex-direction:column;}
.card-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(320px,1fr));gap:12px;}

/* Licitacion card */
.lic-card{background:#fff;border:1px solid #EBEBEB;border-radius:10px;padding:0;cursor:pointer;transition:border-color .15s,box-shadow .15s;overflow:hidden;}
.lic-card:hover{border-color:#C7C7C7;box-shadow:0 2px 12px rgba(0,0,0,.07);}
.lc-top{padding:12px 14px 10px;border-bottom:1px solid #F5F5F5;}
.lc-due{display:flex;align-items:center;justify-content:space-between;margin-bottom:10px;}
.due-chip{display:inline-flex;align-items:center;gap:4px;font-size:11px;font-weight:500;padding:3px 8px;border-radius:5px;}
.due-chip.urgente{background:#FEF2F2;color:#DC2626;}
.due-chip.pronto{background:#FFFBEB;color:#D97706;}
.due-chip.ok{background:#F0FDF4;color:#059669;}
.due-chip.lejano{background:#F5F5F5;color:#888;}
.lc-menu{color:#CCCCCC;font-size:16px;cursor:pointer;padding:2px;}
.lc-menu:hover{color:#888;}
.lc-title{font-size:13px;font-weight:600;color:#111;line-height:1.45;margin-bottom:4px;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden;}
.lc-organo{font-size:12px;color:#888;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;}
.lc-dates{font-size:11px;color:#AAA;margin-top:5px;display:flex;gap:6px;flex-wrap:wrap;}
.lc-date-sep{color:#DDD;}
.lc-bottom{padding:10px 14px;display:flex;align-items:center;justify-content:space-between;gap:8px;}
.lc-cpvs{display:flex;gap:5px;flex-wrap:wrap;flex:1;min-width:0;}
.cpv-chip{background:#F5F4FF;color:#4A4AE8;font-size:10px;font-weight:500;padding:2px 7px;border-radius:4px;white-space:nowrap;max-width:130px;overflow:hidden;text-overflow:ellipsis;}
.lc-importe{font-size:13px;font-weight:600;color:#111;white-space:nowrap;flex-shrink:0;}
.lc-importe.nd{color:#CCCCCC;font-weight:400;font-size:12px;}
.lc-actions{padding:0 14px 12px;display:flex;gap:7px;}
.btn-analizar{flex:1;height:30px;background:#111;color:#fff;border:none;border-radius:7px;font-size:12px;font-weight:500;cursor:pointer;display:flex;align-items:center;justify-content:center;gap:5px;transition:background .15s;}
.btn-analizar:hover{background:#333;}
.btn-analizar .a-spin{display:none;width:11px;height:11px;border:2px solid rgba(255,255,255,.3);border-top-color:#fff;border-radius:50%;animation:spin .6s linear infinite;}
.btn-analizar.loading .a-spin{display:block;}
.btn-analizar.loading .a-label{display:none;}
.btn-ver{height:30px;background:#fff;color:#444;border:1px solid #E0E0E0;border-radius:7px;font-size:12px;padding:0 10px;cursor:pointer;display:flex;align-items:center;gap:4px;white-space:nowrap;transition:background .12s;}
.btn-ver:hover{background:#F5F5F5;}

/* Empty state */
.empty-state{display:flex;flex-direction:column;align-items:center;justify-content:center;flex:1;text-align:center;color:#BBBBBB;}
.empty-state .e-icon{font-size:40px;margin-bottom:12px;}
.empty-state p{font-size:14px;}

/* Modal */
.modal-bg{display:none;position:fixed;inset:0;background:rgba(0,0,0,.4);z-index:100;align-items:center;justify-content:center;padding:20px;}
.modal-analisis{max-width:720px;}
.analysis-body{font-size:13px;color:#333;line-height:1.8;}
.analysis-section{border:1px solid #F0F0F0;border-radius:10px;padding:14px 16px;margin-bottom:12px;background:#FAFAFA;}
.analysis-section:first-child{margin-top:0;}
.analysis-section h3{font-size:13px;font-weight:700;color:#111;margin:0 0 8px;display:flex;align-items:center;gap:7px;border-bottom:1px solid #EBEBEB;padding-bottom:8px;}
.analysis-section ul{padding-left:18px;margin:0;}
.analysis-section li{margin-bottom:4px;}
.analysis-section p{margin:0 0 6px;}
.analysis-section p:last-child{margin-bottom:0;}
.analysis-body{display:flex;flex-direction:column;gap:0;}
.analysis-loading{display:flex;flex-direction:column;align-items:center;justify-content:center;padding:50px 20px;gap:14px;color:#999;}
.analysis-loading .big-spin{width:32px;height:32px;border:3px solid #F0EFFD;border-top-color:#4A4AE8;border-radius:50%;animation:spin .8s linear infinite;}
.no-api-warn{background:#FFFBEB;border:1px solid #FDE68A;border-radius:8px;padding:12px 14px;font-size:12px;color:#92400E;line-height:1.5;}
.no-api-warn code{background:#FEF3C7;padding:1px 5px;border-radius:3px;font-family:monospace;}
.modal-bg.open{display:flex;}
.modal{background:#fff;border-radius:14px;width:100%;max-width:600px;max-height:85vh;overflow-y:auto;box-shadow:0 20px 60px rgba(0,0,0,.2);}
.modal-head{padding:20px 22px 16px;border-bottom:1px solid #F0F0F0;display:flex;align-items:flex-start;gap:12px;}
.modal-head-icon{width:36px;height:36px;background:#F0EFFD;border-radius:8px;display:flex;align-items:center;justify-content:center;color:#4A4AE8;font-size:18px;flex-shrink:0;margin-top:2px;}
.modal-head h2{font-size:15px;font-weight:600;color:#111;line-height:1.4;flex:1;}
.modal-close{width:28px;height:28px;border-radius:6px;border:1px solid #EBEBEB;background:#fff;cursor:pointer;color:#888;font-size:16px;display:flex;align-items:center;justify-content:center;flex-shrink:0;}
.modal-close:hover{background:#F5F5F5;}
.modal-body{padding:18px 22px;}
.mf{margin-bottom:14px;}
.mf label{font-size:11px;font-weight:600;color:#AAAAAA;text-transform:uppercase;letter-spacing:.6px;display:block;margin-bottom:4px;}
.mf p{font-size:13px;color:#222;}
.mf .mf-val{font-size:18px;font-weight:600;}
.modal-footer{padding:14px 22px;border-top:1px solid #F0F0F0;}
.btn-link{display:inline-flex;align-items:center;gap:6px;background:#111;color:#fff;text-decoration:none;padding:9px 18px;border-radius:8px;font-size:13px;font-weight:500;}
.btn-link:hover{background:#333;}

/* Log modal */
.log-pre{background:#111;color:#a8ff78;font-family:'Courier New',monospace;font-size:11px;padding:14px;border-radius:8px;white-space:pre-wrap;max-height:380px;overflow-y:auto;margin-top:12px;}

/* Toast */
.toast{position:fixed;bottom:24px;right:24px;background:#fff;border:1px solid #EBEBEB;border-radius:10px;padding:12px 16px;font-size:13px;color:#333;box-shadow:0 4px 20px rgba(0,0,0,.12);z-index:200;transform:translateY(60px);opacity:0;transition:all .25s;max-width:320px;}
.toast.show{transform:translateY(0);opacity:1;}
.toast.ok{border-left:3px solid #059669;}
.toast.err{border-left:3px solid #DC2626;}
</style>
</head>
<body>
<div class="app">

  <!-- Sidebar -->
  <div class="sidebar">
    <div class="sidebar-head">
      <div class="sidebar-logo">
        <div class="sidebar-logo-icon">
          <svg width="26" height="26" viewBox="0 0 26 26" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="26" height="26" rx="7" fill="#111"/>
            <circle cx="13" cy="11" r="4.5" stroke="white" stroke-width="1.6" fill="none"/>
            <circle cx="13" cy="11" r="1.5" fill="white"/>
            <line x1="13" y1="16" x2="13" y2="20" stroke="white" stroke-width="1.6" stroke-linecap="round"/>
            <line x1="9.5" y1="20" x2="16.5" y2="20" stroke="white" stroke-width="1.6" stroke-linecap="round"/>
            <path d="M7 11 A6 6 0 0 1 19 11" stroke="#4A4AE8" stroke-width="1.4" stroke-linecap="round" fill="none"/>
            <path d="M4.5 11 A8.5 8.5 0 0 1 21.5 11" stroke="#4A4AE8" stroke-width="1.1" stroke-linecap="round" fill="none" opacity="0.45"/>
          </svg>
        </div>
        Centinela
      </div>
    </div>

    <div style="padding:14px 12px 12px;border-bottom:1px solid #F0F0F0;display:flex;justify-content:center;">
      <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAABH4AAAFaCAYAAABosJqeAAAQAElEQVR4AeydCdxN1ffG1zFPRaFCKUmGRIZISTKEUBRJGcsUMmUeX/OQeRaJQqhESiEhFBJJ/kXSQKIoyjzUv+fk/uLtHe5whr33efp03ffee87ea33XucNZZz1rp/iL/5EACZAACZAACZAACZAACZAACZAACZhOgP4FlEAK4X8kQAIkQAIkQAIkQAIkQAIkQAIBIkBXSYAEgkSAiZ8gRZu+kgAJkAAJkAAJkAAJkMClBPg3CZAACZCA8QSY+DE+xHSQBEiABEiABEiABJInwC1IgARIgARIgATMJMDEj5lxpVckQAIkQAIkEC0B7kcCJEACJEACJEACJGAQASZ+DAomXSEBEiABZwlwNBIgARIgARIgARIgARIgAd0JMPGjewRpPwl4QYBzkAAJkAAJkAAJkAAJkAAJkAAJaEmAiR8tw+af0ZyZBEiABEiABEiABEiABEiABEiABEhAHwLRJn708ZCWkgAJkAAJkAAJkAAJkAAJkAAJkAAJREuA+2lOgIkfzQNI86Mj8Pnnn0vlypXl/vvv5y0GBg8++KA8+eST8uyzz0pcXJyMHz9e5syZI8uWLZONGzfK119/LUeOHJG//vorukBxLxIIk8Dp06flwIEDsmPHDlm7dq28+eab8uKLL8rzzz8v3bt3lxYtWkidOnWkQoUKfM9H+Z6fNm2anDlzJsyI6LHZ9OnTpWrVqlofE++8846cPXtWD+C0UjsCv/zyi3z22WeyYsUKWbBggUydOlWGDh0qXbt2lebNm9ufqxUrVtT2PXTo0CHtYuK/wc5bcPToUfnyyy9l9erV8sYbb8iMGTPs7++ePXvKM888I48//rhUqVJF2+MMv0N+/fVX58FxRBKIgAATPxHA4qbmEChSpIg8/PDDsn79elmzZg1vUTJ49913Zd68eTJx4kTp37+/tG/fXho2bCjVq1eXMmXKyK233irZsmWTFClS2H8/9NBD0qVLF/sLfd26dfLzzz8zKWTO28oTT5Dc+eCDD2TSpEnSpk0bO5GTI0cOyZAhg+TKlUvw3i5fvrw88sgj0qxZM/vkZPjw4YITfPyYxI/KNVEe70HeD4yR5E2bNq0ncfZikm+++cZOCi5fvlx0ji2TPl4cLebOcezYMcH3MS7cdOzYUerWrSv33HOP5MmTR9KkSSPXXHONFCtWzD7pxsk3TsJxMo6k+oy/T87xuYrPZF3fQ7hoYG501fHs/PnzdgJx1qxZ9udugwYN7O/vAgUKyBVXXCFXX321FCpUyH4OF2mQVERyEUlGJBuRdETycY2m39+nTp1SJxi0JLAEmPgJbOjpeNu2be2TQyQlSMN9Aqj+Wbp0qYwcOdK+SliuXDm59tprJXv27HL33XdL06ZN7YohnIyxQsj9eOgww+7du2Xy5Ml2AgcnIvhheP311wuuLuP9i9eQyDl48CATiC4GNHfu3IKrlUiuuTiN50Pj5PX48eOez2vihPRJDwJIcqxcuVKGDRsm9erVsy/IZM2aVfB9jAs3Y8eOlddff10++ugj+e677+TcuXN6OEYrlSOAqm9coMEFmBIlSkjGjBntBCJ+6+FizNy5c+3qnl27dgk+h/m7T7kQ0iADCTDxY2BQ6VL4BPr06WMnHsLfg1s6TQBSsI8//lhwFQg/PG+55Ra7agMnZatWrZILFy44PSXHU5QATjTmz58vrVq1kptvvlkKFixoV/VAsoUTkd9++40JHh9ihyq9/Pnz25V7yUyvzcuQrUGSymoZbUJGQ6Mk8NVXX9lVtrVr15arrrpKHnjgAenRo4csXLjQlmPzOzZKsNztMgJ79+4VJHMaN25sV4nhgh4u0OD7e+vWrZSjXkaLD0jAHwJM/PjDnbMqQqBw4cLSr18/SZUqlSIW0QwQ+OKLL+weApUqVZKcOXNKo0aN7H4trAYCHR1u4dmIfjGLFy+W5557zpYGIulXv359wUn5t99+K3/++Wd4A3Er1wggHk2aNBGTJF6oPkRiGVeZXQPHgUnARwJIlOMYh+Qa8lfIZvBZi4ofH83i1IYR2LRpk/Tu3VvuvfdeW6YF+dbLL78s6AvFCh7Dgk13jCDAxI8RYaQTsRCAXv3RRx816mp2LDwc3deBwdAH6JVXXrHlYfny5bPL0xctWsRKIAfY+jUEkjqQGpQsWdJuDDp69Gi7GTivPPsVkYTnveGGG6Rbt252/6SEt9Dz2V69etnSAj2tp9UkkDCBn376SV544QW7Fw+a2KM3CuQ2lGslzIvPRkcACfNXX31VatWqZffjGTx4sN0vExdyohuRe5EACXhFgIkfr0gHfB7V3e/bt6/db0Z1O4NsH64e4fbaa6/ZyYK77rpLRo0aJfv37xc8H2Q2uvi+YcMGQek39P64Go3KLiZ71I0eGmsWKFDAqKQ4moSiKT0lXuoed7QsMgJorIzP1eLFi9urH6EBLk/CI2PIrZMngO/ruLg4KVWqlF2FvWTJEjl58mTyO3ILEiABZQh4nfhRxnEaQgKXEsBKAkj+UPJ1KRV1/0aiZ8uWLfYKYVhtBD96P/zwQ1YBKRoyJOtwdfD++++3V+Nirx5FA3WJWVi9p4lhEi80EWW1zyVB5p9aE0APvCeeeEJq1Khhf66iyT3lsVqHVEnjkfBBnzcssIDVW7HkOlboUtJYGqUqAdqlCAEmfhQJBM3wnwAayj722GNGXd32n6q7FiABdPjwYXvlp/Lly9ulx+vWrWNvGHexhz06rggi2YOTE/xNyUHY6HzdEBKv7t27GyfxQpXZH3/84StbTk4CsRIIJXxq1qwpkNxwmehYiXL/hAhcmvDBaqy///57QptF+Bw3JwES8JMAEz9+0ufcyhHAKl85cuQQy7KUs40GJU0ASaC3335bKlasKDjBY1VJ0rzcfBUreKBaBInUNWvWCK8Ouknb+bFNlHhhWWFIYJh8dP544YjeEPjss89siY0RCR9vkHGWKAjgYhoWPSlbtqww4RMFQO5CAgoTYOJH4eDQNO8JoJ8FJF8pU6b0fnLO6AgBnNgNHz5cUGkyb948Jh0coRreIGjEjfdPuXLlBCt7sI9KeNxU2srEVbwgTcBxiaakKrGmLSJkkDwByLfGjRtnN22eM2eOsMIneWbcIjoCWDgD0kE0Bj927Fh0g3AvEiABZQkw8aNsaGiYXwRatGgh6G+RIgXfHn7FwIl5t2/fLg0bNpTGjRsL/saPZyfG5RgJE5gxY4ZAbocVPk6cOMGG2wljUvrZm266ya6Wy5Ahg9d2ujofKgAp8XIVMQd3iQAaN+NEHKvrIbGOylaXpuKwASbw9ddfS+vWre3fS5s2bRJcQAswDrpOAsYS4JmtsaGlY7EQgOQrV65clHzFAlGBffEjGVU/kH9NmzaNP2ZciAlW5UKytE2bNoLKitgSbML/fCSAk8t8+fIZ1edswoQJsnLlSr73fTyuOHV0BFB1Ubt2bcEqdFylKzqG3Ct5AgsXLpTq1avLCy+8IKyKTJ4XtyABnQkw8aNz9Gi7awRuvfVWgTSAq3y5hjjpgR1+9ciRI4LERO/evYXyI+fgbtu2TcqUKSMzZ84kV+ew+jISKuMaNGggadOm9WV+NybduXOnxMXFCSrQ3BifY5KAGwR++eUXQTJ94MCBwoa6bhDmmCECo0aNkpYtWwoqfnARJ/Q870mABMwkwMSPmXE1xis/HWnWrJlgNSJKvvyMgnNzo/pnxIgRdnPM06dPOzdwQEdCDx80f/zkk0+EPxj1Pgjy589PiZfeIaT1hhD49NNPBX22Zs+ezV4+hsRURTew+MWzzz5rJ8aPHj2qoom0iQRIwAUCuiR+XHCdQ5JA8gRQIZI7d25KvpJHpc0WCxYssBs/s0Fm9CHr3LmzfUX65MmT0Q/CPZUhAIkX+vuYlOQeO3asrFq1ihIvZY4yGpIcgSVLlkjdunVl9erVrKBMDhZfj5oAVodDdef06dMp7YqaInf0iACncZgAEz8OA+VwZhG45ZZbBP1+KPkyK64bN26UvHnzUgISRVhr1aolWGGGPSeigKfgLqhsfOyxxyRNmjQKWhedSZ9//rkMGDCA7+/o8HEvHwigvwrklt9++62wT5oPAQjIlG+//bbUq1dP3nvvPdHrOzwgAaKbJOAyASZ+XAbM4fUn8NRTTwmujqRMmVJ/Z+jB/wj89NNPki1bNl7x+h+R5P9AnyQ0Gj1//nzyG3ML5QkULlxYevXqJaat4gWf2KRU+cOPBl4kMGbMGOnSpYtw+eyLQJK642tRE5g7d640bdpUdu/ezeRi1BS5IwnoTYCJH73jR+s9IoCqH0ghLMvyaEZO4wUB9PopUqSIULKUPO3BgweziXPymLTaAhKvnDlzGiVlxUk0pDJcjlirQzFiY03ZAcdrXFwcmzibElBF/Zg0aZK0a9dODh8+rKiFNIsESMALAkz8eEGZc2hPIE+ePJR8aR/FhB1AaX3lypXZSDNhPPazw4cPl0GDBgkSZfYT/Ed7Aq1btxYsFa25xOuyOKB3BVZC4ipel2HhA0UJQHaDz1Wu3KVogAwxa/369fb396+//mqIR3SDBEggWgJM/ERLjvsFjgD0940aNRJKvswL/UcffWTL+ZjY+G9skfTBFWm12fzXbj6TOIHixYtT4pU4Hr5CAq4T+PjjjwUycp6Mu4460BPs2bPHXiXu4MGDgeZA50mABP4hwMTPPxz4LwmERQCrfKH6x7Io+QoLmNcbxTDfokWL5Mknn2RVyyUM0XC0f//+ZHIJExP+hMQre/bsRkm8Ro4cKWvXruUqXiYcoIb7gObjaJL/yy+/GO4p3fOTwNGjR6V9+/by888/+2kG5yYBElCIABM/CgWDpjhHwK2R0OcH/X64ypdbhP0dF8mfDh06cLWLv8Pw/vvvS6dOnSiB+5uFSf/j+K5Ro4akTp3aGLc+/fRTGTJkCFfxMiai5jqyb98+adu2rbDSx9wYq+IZmtwjGX727FlVTKIdJEACPhMwPfHjM15ObyIByL2aNGlCyZeJwf3bp2nTpsnLL78sQV65au/evdK5c2cmff4+Hkz6v3Tp0oKqRRNX8frjjz9MChV9MZQAmjkjURnk7xdDQ6uUWxMnTpRXX32VyXClokJjPCTAqRIhwMRPImD4NAkkRQAnTzfffLNRUomk/A3aa+hps2PHjsAuedq1a1fZuXNnYP039XiHxOvKK680yr0RI0YImpfyRNqosBrpzMKFC2XWrFlcRdLI6KrjFFY1xCqcv/32mzpG+WYJJyYBEriUABM/l9Lg3yQQJoHcuXNLv379hJKvMIFpttmBAwcEyZ8glkgjObB06dJAVzxpdriGZS6SeVWqVDFK4rV582YZNmwYr2qHdQRwIz8JQOKFPlRcwcunKARkWvTzgeyVUsKABJxukkCEBJj4iRAYNyeBEAE0Am7atCklXyEght2/9dZbghWtgpT8ee2112TChAkSJJ8NO2wTdOfee++Vnj17imkSL1ReaPIvpAAAEABJREFUUuIl/C8CAn5tOm7cOLuK8sKFC36ZwHkDQGDs2LGyadMmfocHINZ0kQSiIcDETzTUuA8JXCSARs958+al5OsiD9PuUPWD5ohB+LF+7NgxGT16NBtbG3YQo4kzqrgyZsxolGeo9NmwYUO0lWlGsaAzahNYuXIlJV5qh8gI67Zs2SKzZ88WJsONCCedIAFXCDDx4wpWDhoUAtdff70tCaLky9yIoxnnuXPnzHXwomfwc9u2bQHr63PReYPvkPSpUKGCUbLUjz/+2K7GO3nypMGRo2umEJg8eTJPxk0JpsJ+YFEKXMBR2ESaRgIk4DMBJn58DgCn159A/fr1pVmzZpR86RzKJGx/9913BSt9mZz8wYk0+k+cOXMmCRJ8STcCFStWlC5dukj69Ol1Mz1JeyHxOn78eJLb8EUSUIHAlClTBM12KZ9VIRrm2rBmzRqBVPvEiRPmOknPSIAEYibAxE/MCDmASQSi9QUnIvny5aPkK1qAiu+HapjvvvtO/vrrL8Utjc48Sryi46byXujn0717d+P6+mC1mo0bN1LipfLBR9tsAmjkPHPmTGGS0sbBf1wkgBXjKPFyETCHJgFDCDDxk3Ag+SwJREQgZ86cEhcXZ5ScIiIAhm/8/fff2/1vTLxqO2PGDHnnnXd4Im3YMQyJ1z333GPUZxJ6+qAyjRIvww5WQ91ZtGiR7NmzR4LQI87QEGrh1meffSZLlizh6oZaRItGKk7AePOY+DE+xHTQKwKPPfaYtGjRgpIvr4B7PA/kXpB9mfQj/ocffpBRo0bJ6dOnPabJ6dwkUK1aNenQoYNREi+879BMn9UTbh45HNtJAm+++aYwSekkUY6VEAEkffi5mBCZWJ7jviRgJgEmfsyMK73ygYBlWQLJ16233krJlw/83Z4SMi9Ivkyq+oE/e/fuNVbC5vYxoeL4WbJkkR49ehgn8Ro6dKi9TPH58+dVxE6bSOAyArhIgAo1k74vLnMwKA8U9/Onn36S119/nXJCxeNE80hAFQJM/KgSCdphBIHrrrtO+vfvb5S8wojAOOTEhx9+KEiWmNDoGScmL7zwgvDExKGDQ5FhunbtKnfeeadRn0Fr164VSrwUOcACakakbq9atUpOnToV6W7cngQiIrBixQpB8ufPP/+MaD9uTAIkEEwCTPwEM+702kUCderUkVatWlHy5SJjP4dG4ueLL74Q3X9owQ+u4uXnkeT83A899JC0adNG0qVL5/zgPo2IxGTfvn1VuaLtEwVOqxuBzZs3M6muW9A0tBereVFOqGHgaDIJ+ESAiR+fwHNacwlYliXoRZE/f35KvgwM8+HDh+1GzzpX/aCvz7p169h0NOrjU70dr7nmGunZs6dkzJhRPeNisAgSry1btvBYjYEhd/WWwLfffiu7d+9mw3xvsQduNqzitXr1alaWBS7ydJgEoifAxE/07LgnCSRKIHv27DJgwACj5BaJOhvAF+bOnWvr6s9r2G9kx44dtmyG1T5mHbhYxato0aJGVRp+8MEHtrSSV7TNOlZN9+aTTz5htY/pQVbAPyR9+NmoQCBoAgloRICJH42CRVPVJZCQZY888oi0bt3aqBOxhPwM4nNo9Dx69Ggtr7TB7iNHjrChs0EHLuSlzZs3N0rihZXmKPEy6CANkCvbtm1j4idA8fbLVSQY8Tnp1/yclwRIQD8CTPw4GzOORgL/I2BZ/0i+ChQoQMnX/6iY88fWrVtl+PDhWv3Anz9/vixYsEB0lqmZcwQ548n1119vS7wyZMjgzICKjAKJF95jFy5cUMQimkEC4RHYuXMnP2PDQ8WtYiDwf//3f1r9/ojBVe5KAqoT0MY+Jn60CRUN1ZFA1qxZZeDAgZR86Ri8MGxG9cymTZu06D9y7NgxuzcRJV5hBFajTSDxKlSokFGVhStXrpRx48ZpWVGn0aFDU10iwMSPS2A57GUEmPi5DIciD2gGCahNgIkfteND6wwgUKtWLWnbtq1RJ2YGhMURF7BcL1bHwspDjgzo4iAjR46Uzz//XPvVyFxEpN3Q9evXlyZNmkjatGm1sz0xg0+cOCH9+vXjKl6JAeLzShP45ptvBMcw5MBKG0rj3CXg8ug4zo4fPy48zlwGzeFJwDACTPwYFlC6ox4By/pH8oWr8pZlqWcgLYqJwOLFi+Xll19WurR//fr1dpNcVvvEFGqlds6TJ4/07t1bTJR4ffbZZ1pU0Sl1QNAYJQjs27fvstW8lDCKRhhHYP/+/TzOjIsqHSIB9wkw8eM+Y85AAnLVVVdR8mXocYArbqj6OXz4sLJX3yBJY9LHrAMQEq+8efNKihTmfI0vX75cJk6caKLEy6yDj94kSgCSWnwnJLoBXyABBwj8/vvvTI47wJFDkEDQCJjzizFokaO/2hF46KGHpF27dpR8aRe55A3etWuXso2ep06dKu+++y6vDiYfRg+2cGaKJk2aSIMGDYySeOGEmat4OXN8cBT/COCE/M8///TPAM4cCAI8zgIRZjpJAo4TYOLHcaQckAQSJmBZli3NuO2227jKV8KItH52/PjxsmbNmvCuwnnk6Q8//CCjRo0SVvt4BNyDafLnzy+9evWS9OnTezCbd1NghbwdO3Yo9f7xznvOZAoBnpCbEkm1/eBxpnZ8aB0JqEqAiR9VI0O7jCSQJUsWGTRo0P9W+TLSyYA6hfJ+JFlUavT8/PPPC5I/sC2gYTHObUi8cufObZTEa9myZTJp0iRKvIw7WoPnEJLs/LwNXty99hi/M3iceU2d85GA/gSY+FEjhrQiQARq1KghHTp0YPLHwJi///77MnnyZMGPMr/de+edd2TmzJlK2OI3C1Pmb968uTz22GOSJk0aU1yS3377zV7FCyshGeMUHSEBEiABEiABEiCBpAl4/ioTP54j54RBJ2BZlHyZegzgChwaPatQZYPqI1x9NpV10PwqXLiwLRU1bRWvESNGyM6dOynxCtoBTX9JgARIgAQuEuAdCXhDgIkfbzhzFhK4jMCVV14pgwcPltSpU1/2PB/oT+DHH3+0Gz37mXRBv5SPPvqIJ9P6H07/86B79+5y3XXXGdUf7K233rIr5E6dOvU/P/kHCZAACQSWAB0nARIgARcJMPHjIlwOTQJJEXjwwQelY8eOlHwlBUnT11588UVB35Lz58977gEa5I4cOZINnT0n796Ebdq0kdq1axsl8Tp8+LDExcXJ8ePH3QPHkUlAUwI0mwRIgARIgAScJsDEj9NEOR4JhEnAsv6RfN1+++1GXcUP032jN4PkC1IrP3r9QDqDpbGNBhwg50qUKGFLvExbxQvH6Zdffilc+jrJg5kvkgAJkAAJkAAJkIAjBJj4cQQjByGB6AhkypSJkq/o0Cm/F6RWXid/5s2bJ2+88YacO3dOeT40MDwCWMUra9asRiWHFy9eLNOmTZPTp0+HB4FbkQAJkAAJkAAJkAAJxESAiZ+Y8HFnEoidQNWqVeW5556j5Ct2lMqNMHr0aHG0qiEJD1HlQ4lXEoA0fAmr/1WvXt2oXmCHDh2ixEvDY5EmkwAJkAAJkAAJ6E2AiR+940frDSBgWf9IvooUKRL2VX0D3A6EC0ePHpVhw4aJF5IvNHTG6kiUzphxaN11113Sp08fMW0Vr+eff1527dpFiZcZhym9IAESIAESIAES0IQAEz+aBCoRM/m0IQRwcjdkyBCjruwbEpqY3ViwYIEtv3Kz0fP69etl/PjxniSYYgbCAcIi0L17d7niiivC2laXjSBDnD59OiVeugSMdpIACZAACZAACahGIGp7mPiJGh13JAFnCTzwwAPSpUsXSr6cxer7aGj0jCoHN/uZoFHumTNnfPeVBjhDAH198HmQOnVqZwZUYJQDBw5Q4qVAHGgCCZAACZCAKQToBwlERoCJn8h4cWsScI2AZVm2tOOOO+6QFCn41nQNtA8Db9++3W7i7UZyZsqUKbJy5Upxs6LIB2SBnbJcuXLSs2dPMXEVrz179lDiFdgjm46TAAm4RoADkwAJkEAYBHh2GQYkbkICXhFImzatQPKVKlUqr6bkPB4RwApfn376qaMnvj/88IOw2sejAHowDSp8UO0D6acH03k2xcKFC+Wll16ixMsz4pwoqAToNwmQAAmQAAkkRoCJn8TI8HkS8IlApUqVBCd/TP74FACXpj137pyg0bOTVT9o6Pzjjz8K5GQumc1hPSSA9/39999vlNxz37590r9/fzl+/LiHJAM/FQGQAAmQAAmQAAmQwGUEmPi5DAcfkID/BCzrH8lXsWLFKPnyPxyOWvD222/LK6+8IkgCxTrwO++8I7Nnz3ZkrFht4f6xE6hYsaLd48tZiVfsdsU6Avpb7d2719FKt1ht4v4kQAIkQAIkQAIkEDQCTPwELeL0VwsCkHygOgT3WhhMI8MigMocnAj/+uuvMVfpQOIVdsPosKzjRn4RyJgxo3Tv3l1Mk3jNnz9fZs2aRYmXXwcW5yUBEiABEiABEiCBiwSY+LkIgnckoBoBSD5wMuik5Es1H4NoDxrcoo/T2bNno3YfEq/NmzfLhQsXoh6DO6pDABKve+65xyiJ13fffWdLvE6cOKEOaFpCAiRAAiRAAiRAAgElwMRPMANPrzUgYFn/SL6KFy9OyZcG8YrExAkTJsi6deuiStzs2LFDUA3Gap9IiKu7bbVq1aR9+/bGreKFyrZvv/2WEi91Dz1aRgIkQAIkQAIkEBwCwsRPgIJNV/UjkDJlSkF1ByVf+sUuKYsh+Ro6dKhE0+gZSR82yk2Krj6vZcmSRXr06CGQeuljdfKWzp071+5lFc3xnfzo3IIESIAESIAESCB6AtwzqASY+Alq5Om3NgTuu+8+6dWrl1EyEG3gu2jo6tWrZfr06RE1Z543b568+eabcv78eRct49BeEYDE68477xQkeL2a0+15vvnmG0q83IbM8UmABEjACQIcgwRIIFAEUgTKWzpLAhoSsCzLTvyUKFGCki8N45eYyaj6QfXOvn37wmr0fOzYMbv6i1UUiRHV6/mHH35Y2rRpI+nSpdPL8GSshcTrhx9+oMQrGU58mQRUIkBbSIAESIAEzCfAxI/5MaaHBhBIkSKFfdJPyZcBwbzEhYMHDwoaPYeTzIHk76uvvuIJ9SX8dP3zmmuusSVepq3iNXv2bEFVWjjHs66xM9xuukcCJEACJEACJGAoASZ+DA0s3TKPQLly5aRPnz6UfIlZ/82cOVOWL1+epHxr/fr1Mn78eIllJTCzqOntDSReRYsWVVjiFTnfr7/+WgYOHCjsPxU5O+5BAiRAAiRAAiRAAm4TYOLHbcIcnwQcImBZlvTs2VPQEwQVQA4Ny2F8JhCSfCWV1EG1jy+rePnMxsTp69SpI82bNzdS4rV///6wZIsmxpU+kQAJkAAJkAAJkIDKBJj4UTk6tI0E4hGwLMsXyVc8M/jQYQKbNm2SsWPHJljRM2XKFFm1apVcuHDB4Vk5nPX+iOkAABAASURBVNcErr/+ejt5a9oqXi+99JLMnz8/qlXqvI4B5yMBEiABEiABEiCBIBJg4ieIUY/eZ+6pAIGyZctKv379hP1+FAiGQyaEqn527dp1WQ8fNMlFDyBW+zgE2udhIPEqVKiQUU3a0XdqwIABlHj5fGxxehIgARIgARIgARJIikCUiZ+khuRrJEACbhKwLEu6d+9OyZebkH0Y+48//pChQ4deVvUDidehQ4cEiSEfTOKUDhJ44oknpEmTJpI2bVoHR/V/qBEjRshPP/3EY9T/UNACEiABEiABEnCRAIfWnQATP7pHkPYHkoBlWYITLlb9mBV+yGXefPNNu9HzO++8I7NmzZJz586Z5WQAvbn55pttiZdpq3jNmDFDXnvtNUq8AnhM02USIIEAE6DrJEACWhJg4kfLsP3X6JMnT8rPP/8se/fule3btwtWAXrvvffsH+VYNQgrAkEygpL8aG9//vnnfyfmM74RuPvuu6V///6UfPkWAecnRmUPqn7OnDkjw4YNE0q8nGfsx4iQeN1yyy1GSbx27txpr+J14sQJP5ByThIgAQUI0AQSIAESIAF9CDDxo0+sbEu///57effdd2XUqFHy9NNPS5kyZSRLlixy5ZVXSs6cOeXWW2+V4sWLy3333SfVq1eXxx9/3F5BpmPHjtK7d2+Ji4uL+IaqA8gULMuybeA/ahCwLEu6du0qpUqVMuqEUg26/lmxY8cOW/L1ySefXNbvxz+LOHMsBJo0aSL4/DRR4nXw4EFKvGI5OMzZl56QAAmQAAmQAAkoToCJH4UD9PXXX8vEiROlWbNm/0vw5M2bV2rUqGGf8GMlFawGdOzYMXvFH6z6gxsqcxK6oZogmhsSRpAqWBYTP6odLpZlyfPPP8+qH9UCE6M9qM5D1U+Mw3B3nwkUKFBAevXqJaZJvKZNmyZvvPHGZf2o/kHNf0mABEiABEiABEiABFQkwMSPQlHB1dNFixZJu3btJH/+/FKwYEFp3769QKp1aYInlNQJJXHcdKFp06bSqFEjVpS4CTnGse+66y5bcsF+PzGCVGh3vLcVMidyU7iHTQASr9y5cxv1+Qkp8cCBAwXyYttJ/kMCJEACJEACJEACJKA8ASZ+fA7RypUrBUtzV65cWW644QapW7euXeWze/duu4oHSR6cBOLmtan58uWzbUuZMqXXU3O+CAhYliWdO3dWUvIVgRvclASMItCiRQv78zxNmjRG+YUKw19++YUSL6OiSmdIgARIgARIgARMJ8DEjw8RxvLMkHDde++9Uq1aNbtaY9WqVfZKPqFEjw9m/WfKPn362Mkoy4pZ4vWfsfmEswQsy5LRo0dT8uUsVo5GAlERuP322+2eaqZJvCZPniyLFy+mxCuqo4I7kQAJkAAJkAAJkIB/BDxO/PjnqAozb9iwQTp06CBFixa1JVx4jJ48qObBTQUbQzbganX9+vWNkiiEfDP1/s4775TBgwcz+WNqgOmXNgS6d+8u1113nViWOUnzbdu22RcpKPHS5jCkoSRAAiRAAiSgAAGaoAoBJn48iMTLL79sN2TGSlvjx48XVPyoVNkTHwF6C/Xt21dSpUoV/yU+VpiAZVnSqVMnQc+fFCn41lY4VDTNYAJt2rSRWrVqGZeAHTFihBw5coQSL4OPXbpGAiRAAq4S4OAkQAK+EuDZoYv4FyxYIEj2oEHysmXL7J49qlX2JOQ+JF45cuRI6CU+pzgBy6LkS/EQ0TyDCZQoUcKWeKVPn94oL3HBYunSpXLu3Dmj/KIzJEAC/hDgrCRAAiRAAt4TYOLHBebLly+XmjVrypNPPinr1q0Tlat74rvfunVruyEpK0bik9HnMU4+hwwZYlzFgT4RoKVBJYBVvLJmzWqUxGvLli0yaNAgOXHiRFDDSr/dI8CRSYAESIAESIAEPCLAxI+DoDdu3CiNGjWyZV3vvPOONhU+IQRoSEqJV4iGvveWZQl6SZUpU4Y9mvQNIy3XjADec9WrVzcu4YpVvI4ePepyNDg8CZAACZAACZAACZCAmwSY+HGA7vnz5wXNPLFK15w5c+zVuXSQdMV3HUmf7Nmzx3+ajzUkgIotrvKlYeCCbrKm/qOvVu/evcW0VbzGjh0ruIhBiZemBybNJgESIAESIAESIIGLBJj4uQgi2rvNmzdLtWrVBFdFkQDSMeED39u1a2c3JEXCAI95059A8eLFZfjw4VpWIOhPnx4EiUDXrl0lc+bMRrm8adMmSryMiiidIQESIAESIAESCDIBJn5iiP60adOkatWqsmrVKruPTwxD+bprsWLFBNU+Cq7i5SsX3Se3LEueffZZufvuuyn50j2YtF9ZAl26dJEqVaoYtwoiLmb8/vvvynKnYSRAAiRAAiRAAiRAAuET0CTxE75DXmx5+PBhadmypbRp00Z+++030bXKJ8QKq3hdddVVoYe8N4gAKrjGjBnDqh+DYkpX1CEAeW+vXr2Mk3iNGjVK3nvvPa7ipc6hRktIgARIgARIIGAE6K7TBJj4iZAoVjiBtGvGjBl28+YId1du806dOtkrkCFBoJxxNMgRAnfccYeMGDGCyR9HaHIQEviXAFbxypgx479PGPDXxx9/LIMHD+YqXgbEki6QAAmQgBEE6AQJkIAjBJj4iQDjhx9+KHXq1JFPP/1Ua2lXyOVSpUpJv379jJMohPzj/T8ELMuStm3byj333EPJ1z9I+C8JxEwAlT4VKlQw7vNz/fr1cvr06Zj5cAASIAEScJoAxyMBEiABEoieABM/YbJD0qd69ery/fffi+7SrpDLkHhdccUVoYe8N5gAKrqwQk/q1KkN9pKukYA3BJDwQUPn9OnTezOhh7OY8v3mITJO5T0BzkgCJEACJEACJBAhASZ+wgCGpE+lSpXk+PHjYWytxyaQKKAxtWVZehhMK2MmUKRIEXv1OSZ/YkbJAQJOoHv37sb19dEzpLSaBEiABEiABEiABEggHAJM/CRDKZT0OXfuXDJb6vMyJD9cxUufeDllqWVZdkNyNKRFBZBT43IcEvCdgIcGxMXFSdmyZY2TeHmIkFORAAmQAAmQAAmQAAl4TICJnySAnzp1SipXrmzcyia9e/cWEyUKSYSSL10kgISPyZKvi27yjgRcIYAqyY4dO/Lz0xW6HJQESIAESIAESIAESMAtAkz8JEEWPX1MqvSBq2hICtmaZRkt8YKrvCVCoHDhwjJy5Eiu8pUIHz5NAgkRQNIUElnTVvFKyFc+RwIkQAIkQAIkQAIkYBYBwxM/0Qdr1KhRsnbtWmMaOYPEfffdJ5R4SeD/syxLWrduLeXKleMqX4E/GgggXAIDBw6Uu+66S1KmTBnuLtyOBEiABEiABEiABEjAUwKcLDECTPwkQGbx4sWC5p1//vlnAq/q+xRW8WJjX33j56TlqF4YN26cpEmTxslhORYJGEmgRo0a8uyzz0q6dOmM9I9OkQAJkAAJkIBxBOgQCZDAZQSY+LkMh8iOHTukQ4cOcv78+Xiv6P2wX79+Ur58ebEsSrz0jqRz1hcqVEhQ2cbkj3NMOZJ5BPD+wIWADBkymOccPSIBEiCBABCgiyRAAiRAAiJM/MQ7Cvr37y8//PBDvGf1flixYkXp2bMnJQp6h9Fx6y3LklatWgkkgKgAcnwCDkgCBhAYPHiwlChRgp+fBsSSLgSeAAGQAAmQAAmQQGAJMPFzSeix2hFkXn/99dclz+r9p2VZQomX3jF003okfMaPH0/Jl5uQOba2BB555BFp2bIlJV7aRjAxw/k8CZAACZAACZAACQSLABM/F+O9ZcsWQbXPhQsXLj5jxt2AAQOkbNmylHiZEU5XvMifP7+MHj2ayR9X6HJQpQkkYRykXZB4cRWvJCDxJRIgARIgARIgARIgAS0IMPFzMUxI+hw7duziIzPuqlSpIlh+mKvQmBFPt7ywLMuuagiy5MstthxXXwKQeBUpUoQr3+kbQlpOAiRAAiRAAiRAAiRwkQATP3+DeP7552XZsmViksQLDUm5dPvfwY3s/8BuDcnXxIkTWfUT2COAjl9KoF69etKsWTNJmzbtpU/zbxIgARIgARIgARIgARLQkkDgEz8ff/yxxMXFyeVLt2sZy8uMHjhwoJQuXZoSr8uo8EFSBPLlyydjxoxh8icpSHzNeAJXXnml9OjRQyD1Mt5ZOkgCJEACJEACJEACJCAi5kMIfOIHEq9Tp04ZFemaNWtKx44duQqNUVF13xnLsqRFixZSvnx5ylvcx80ZFCUwdOhQKVCgAN8DisaHZpEACZAACZCAqwQ4OAkYSiDQiZ8hQ4bIypUrjZJ44So1VvFKlSqVoYcs3XKTACRfkyZNosTFTcgcW1kCDRs2lEaNGvH4VzZCNIwESIAEvCPAmUiABEjAJAKBTfx8+OGHYqrEq0SJEpR4mfQu9diXvHnzUvLlMXNO5z+Bq666SrCKF5Ln/ltDC0iABBQiQFNIgARIgARIQHsCgUz8oIkzJF7nz5/XPoCXOlCrVi1p164dJQqXQuHfEROwLEuaN28uFSpUoFwwYnrcQVcCw4cPl1tuuYWfn7oG0BO7OQkJkAAJkAAJkAAJ6EkgkImfAQMGyJo1a4ySeKEhab9+/Xiiruf7UDmrQ5IvrA6nnHE0iAQcJtC0aVN54oknwm9s7vD8HI4ESIAESIAESIAESIAE3CQQuMTPBx98IFjxysRVvIoUKUKJl5vvloCNnSdPHhk3bhxPhpOIO1/Sn0D27NmlZ8+eXMVL/1DSAxIgARIgARIgARIggUQIBCrxc/bsWYHE68KFC4ng0PPpunXrSuvWrSlR8C98Rs5sWZY8/fTTUrFiRVaSGRlhOgUCzz//vOTOnZtJc8DgjQRIgARIgARIgARIwEgCgUr8oNJn/fr1LgbS+6Gvvvpq6du3L0/MvUcfiBkp+QpEmAPrJHpZ1alTh1VtgT0C6DgJkAAJkAAJkAAJxEJAn30Dk/h57733ZNiwYWKaxGvQoEFSqFAhXq3W5z2nnaU33XQTJV/aRY0GJ0cgZ86c0qtXL0q8kgPF10mABEiABEiABJInwC1IQHECgUj8nDhxQtDQ2bRVvOrXry/NmjWjxEvxN5nu5lmWZUu+KlWqxMoy3YNJ+/9HABIvJH8sy/rfc/yDBEiABEiABGIlwP1JgARIQEUCgUj8IOmzadMmFflHbdM111xjS7xSpUoV9RjckQTCJQDJ1+TJkymJCRcYt1OawDPPPCO1atWS1KlTK20njSMBEtCaAI0nARIgARIgAWUIGJ/4Wbp0qYwePdpIidett95KiZcybyXzDUED3PHjxzP5Y36ojfbwhhtukD59+kj69OmN9pPOqUSAtpAACZAACZAACZCAvwSMTvz89ttv9ipepkm8GjVqJE2bNqXEy9/3TuBmtyxLnnrqKalcuTIlX4GLvjkOjxo1SrJly+ZP0twcjPSEBEiABEiABEiABEhAIwJGJ36wite2bds0CkdcvWv7AAAQAElEQVTypqInBVbxosQreVbcwnkCkHxNnTqVVT8xouXu/hBo166dVK9enRIvf/BzVhIgARIgARIgARIgAZ8IGJv4WbRokUCWYtoqXoMHD5Y8efL4dLhwWocJaDlcrly5ZMKECUz+aBm94BqNz01IvDJkyBBcCPScBEiABEiABEiABEggkASMTPz88ssvtsTrwoULmgQ1PDMh72rQoAElXuHh4lYuEbAsy5YaVqlShZIvlxhzWOcJjBkzRrJkyeL8wByRBEiABEiABEiABEiABCIi4P3GRiZ+sIrXF1984T1NF2e88cYbpV+/fkKJl4uQOXTYBCj5ChsVN1SAQKdOneSBBx7g56cCsaAJJEACJEACJEAClxDgnyTgEQHjEj8LFiyQKVOmGLmKF1aj8ei44DQkkCwB9JuaNGkSJV/JkuIGfhLIly8fV/HyMwCcmwRIgARIICwC3IgESIAE3CRgVOLnwIEDRkq8WrRoIY8//jglXm6+Ezh2VAQaN24sVatWpeQrKnrcyQsCY8eOlUyZMnkxFecgARIgAScIcAwSIAESIAEScJyAUYkfSLy++uorxyH5OWDevHkp8fIzAJw7SQKQfE2bNo1VP0lS4ot+EejSpYvcf//9lHj5FQDOGyMB7k4CJEACJEACJEACzhAwJvEzd+5cmTFjhvz111/OkFFkFKzidd111yliDc0ggf8SwPE5efJkJn/+i4bP+EigYMGC0rNnT0mfPr2PVjg0NYchARIgARIgARIgARIggRgIGJH4+f777yUuLk5MW8WrdevW8uijj1LiFcMBzl29IdCoUSOpVq0aJV8u4+bw4RMYN26cXHHFFeHvwC1JgARIgARIgARIgARIwFACRiR+IPH65ptvjApR/vz5KfEyKqKOOqPcYJB8vfDCC6z6US4ywTSoR48eUrZsWSYigxl+ek0CJEACJEACJEACJBCPgPaJn1mzZsns2bONlHhly5YtXrjiP+RjElCHwDXXXGOvqJcmTRp1jKIlgSNw++23S7du3SjxClzk6TAJkAAJkAAJkAAJmE4gev+0Tvzs2bPHSIlX+/bt5eGHH6bEK/rjmnv6RKBhw4by4IMPstLCJ/6cVmTChAlcxYsHAgmQAAmQAAmQgNkE6B0JREhA68RP//795YcffojQZbU3v+222yjxUjtEtC4JApB8TZ8+nZKvJBjxJfcI9O3bV0qXLs3Eo3uIOTIJkAAJkIBiBGgOCZAACYRDQNvED/qJvPrqq0ZKvLJkyRJO7LgNCShJABLFqVOnMvmjZHTMNapYsWLSqVMnSZcunblO0jMSIAESSJwAXyEBEiABEiCBRAlomfj58ssvjZR4Pffcc1K9enWxLCvRgPEFEtCBQIMGDexjOWXKlDqYSxsNIDBx4kRKvAyII11wggDHIAESIAESIAESIIHLCWiZ+IHE6+DBg5d7ovmjokWLUuKleQxp/r8EKPn6lwX/cp8AvhOKFy9OiVd81HxMAiRAAiRAAiRAAiRAAn8T0C7xM2nSJHn99deNlHhlypTp75DwfxIwg0DWrFmFki81YmmyFSVLlhQ0xKfEy+Qo0zcSIAESIAESIAESIIFYCGiV+Nm+fbuREq/u3btL1apVKfGK5UjmvuEQ8HwbSL5q1KjBSgzPyQdnwsmTJ8sVV1wRHIfpKQmQAAmQAAmQAAmQAAlESECrxM+AAQPkyJEjEbqo9uYlSpSwJV7e9kJRmwmtM4cAJF8vvviipE2b1hyn6IkyBAYPHiy333674DhTxigaQgIkQAIkQAIkQAIkQAJKERDRJvEzduxYWbJkiXESryFDhvCkWLE3Bc1xlgBWqcMqfGnSpHF2YI4WaAJ33XWXtG3blqt4BfoooPMkQAIkQAIkQAIREeDGgSWgReLn008/FTTvvHDhglGB6t27t1SsWJESL6OiSmcSIlC/fn2pWbMmJV8JweFzURGYMmUKV/GKihx3IgESIAESIAERMiABEggWAS0SP0j6HDt2zKjI4Gp1nz59eCJsVFTpTGIEIMWh5CsxOnw+UgLDhg2TggULUuIVKThuTwIkQAL/JcBnSIAESIAEAkBA+cTPyJEjZdmyZcZJvNCbInXq1AE4xOgiCfxDIHPmzELJ1z8s+G/0BMqWLSvPPPMMJbLRI+SeJJAIAT5NAiRAAiRAAiRgKgGlEz8bN260Gx+bJvGKi4uT++67jxIvU99V9CtRApB8PfTQQ6x0S5QQX0iOwLRp0yjxSg5SrK9zfxIgARIgARIgARIgAaMIKJ34gcTr1KlTRgHH1eqePXvyxNeoqNKZcAlA8jVz5kw25A0XmM/bqTb9qFGjJG/evJR4qRYY2kMCJEACJEACJEACJKA0AWUTP0OHDpWVK1caJ/HCKl6pUqVS+qCgcSQQj4CjD6+44gpb8pU2bVpHx+VgZhMoX768NGvWjBIvs8NM70iABEiABEiABEiABFwgoGTiZ926ddKvXz8xTeI1aNAgufvuuzWWeLlwBHLIQBJ4/PHHpVatWqx8C2T0o3MaEi8kDaPbm3uRAAmQAAmQAAmQAAmQQHAJRJf4cZHXX3/9JZB4nTt3zsVZvB8aPX26du3KE13v0XNGBQlA8oVVvtKlS6egdTRJNQJjx46VG2+8kUlz1QJDe0iABEiABEiABIJBgF5qT0C5xA+SPqtXr9YebHwHIF2jxCs+FT4OMoGMGTPK9OnTKd0J8kEQhu+VKlWSpk2b8jgJgxU3IQESIAESIAG3CXB8EiABPQkolfj54IMPZPDgwfLnn3/qSTMRq4cPHy6lSpXi1epE+PDp4BKoV68eJV/BDX9YniM5SIlXWKi4EQmQAAl4SYBzkQAJkAAJaERAmcTP2bNnJS4uTs6fP68RvuRNxdXqjh07UuKVPCpuEUACkHy99NJLXOUrgLEPx+WJEydKzpw5mTQPBxa3IQHfCHBiEiABEiABEiAB1Qkok/gZMGCAbNiwQXVeEdtHiVfEyLhDwAikT59e0O+Hq3wFLPDJuFulShVp1KiRpEmTJpkt+bIyBGgICZAACZAACZAACZCAkgSUSPwsX75chg0bZpzEa9SoUVK8eHFerVby0KdRKhGoW7eu1K5dW9gHS6WoRG+LE3siGUiJlxMkOQYJkAAJkAAJkAAJkEDQCfie+Dlx4oQt8bpw4YJRsahatao8++yzAimLUY7RGRIIn0DYW+J9MmvWLDbwDZuY2RtOnTpVsmfPbraT9I4ESIAESIAESIAESIAEPCLge+IHEq/Nmzd75K4306BqARKv1KlTezOh8rPQQBJIngCkXjNnzmTyJ3lURm9Ro0YNeeKJJyjxMjrKdI4ESIAESIAESIAESMBLAt4mfuJ5tnTpUoEcyrRVvOBTkSJF4nnLhyRAAskRqFOnjjzyyCOUfCUHyuDXkfyjxMvgANM1EiABEiABEiCB4BCgp8oQ8C3xc/ToUenfv7+YJvHC1epnnnmGEi9lDnEaohMBSL5efvllVv3oFDQHbX3ooYckQ4YMDo7IoUiABEiABEiABFQgQBtIgAT8JeBb4gcSr61bt/rrvcOzp0uXzm5STYlX9GD37dsna9euFVz17/93YrBLly7SqlUrefLJJ+Xhhx+WihUrSqlSpaRgwYJyww03yFVXXWUnCSATivRWsmRJOXz4cPTGck9XCEAqiQQA7l2ZgIMqS+Dtt9+W1157Tc6dO6esjTSMBEiABEggJgLcmQRIgARIwAcCviR+Fi1aJOPHj5e//vrLB5fdm3L06NF2QsK9GcwYGQ29V6xYIVOmTBEkdiDtueOOOwTyjrx580qlSpWkZcuWMnDgQBkzZoxMnz5d5s+fLzgpXL16tWzZskV27dol+/fvF1SOnT17VqK5de3aVa6++mozoBrkBT4fFi9eLOfPnzfIK7oSDgHIfps3by6///57OJtzGxIgAa0J0HgSIAESIAESIAGvCHie+EGFBSo5TJN41a5dW3DCAqmKV8HTaZ5NmzbJiBEjpHLlypI1a1apXr26veoZEjtLliyRzz//XI4fP25f6ccJP244RnDDyWDohmRh6BaL/1hxrVatWpTkxQLRhX1R6fHUU0/JqVOnXBidQ+pAAO/9Zs2aCRLEOthLGx0gwCFIgARIgARIgARIgARcJeB54gdJnx07drjqlNeDo1Jl2LBhbEh7CfidO3fKjBkzpGnTppI7d2659957pWfPnrJq1So5c+aMXc2BpA5uSOogmSMe/VesWDHp168fVw3yiHck0zRu3NhOAEayD7c1h0DIk7feesuu8kMlX+g53pMACZAACZAACZAACZAACURHwNPEz4IFC2x5j5cn+dFhiWwvrOJ1yy23RLaTgVtDnoHePNWqVZPixYvbvXnQqBd9e1DJgSSPCrHv3bu33RvIwBBo7dKrr74qkHhduHBBaz9ofOwEkAxGby9IOWMfjSOQAAmQAAmQAAmQAAmQQLAJeJb4OXDggJGreNWtW9euagmyxGvbtm3Sq1cvKVGihN2bZ/ny5XbPHSR6cAInUf/n/I4dO3aUmjVrUuLlPNqYRoTMDyf6lHjFhNGonSH5evrppyn5MiqqdIYESIAESIAESIAESMAPAp4lfiDx+uqrr6LzUdG9smTJYq/iFdTVh9544w2pV6+elClTRoYPHy7ffPONLeFSoaonoUPmzjvvtCVeXHUtITr+PgeJF3u6+BsDFWdftmyZzJkzx04kq2gfbSIBEiABEiABEiABEnCBAId0nIAniZ+5c+fKiy++KKomBKKlilW8brrppmh313Y/SHLKlSsn9evXl9dff93u2YPqHtXjC4kX+jFpC95Qw2fPni04wccxZKiLdCtKAqgYbNu2rWBRANU/X6J0kbuRAAmQAAmQAAkkQYAvkQAJOEPA9cTP999/b1dZmHZSh6RHw4YNAyUZwsk5ZFKozli/fr29AhdOzJw5FN0dpXPnzoLeQ0GW5LlLOLrRjxw5Iu3atZPTp09HNwD3Mp5ASPJ18uRJ432lgyRAAiRAAokS4AskQAIkQAIxEHA98QOJ1969e2MwUb1ds2XLZkubgiLx2rBhgzRo0EBq165tV2agUbNOV98hRYuLixNKvNR7LyGJSImXenFRzaIVK1YIKsO4ypdqkaE9JOAHAc5JAiRAAiRAAiQQKQFXEz+zZs0SrOqkU5IgHICQeOXKlSucTbXeBvIKyCwqVqwokHfhpEuXCp9LwaPxdPr06S99in8rQOCFF16QVatWiWnVgAqgNc4EfO60b99efvrpJ+Mkw8YFy0uHOBcJkAAJkAAJkAAJkEBYBFxL/OzZs8dIiRfkXZB5mS4ZgpTr4YcflmnTptk9fHDiFdYRpdhG3bp1kwceeCBQkjzFQpCgOTiB79KlCyVeCdLhkwkRgOSrefPmCa7yldD2fI4ESIAESIAESIAESIAESOAfAq4lfiDx2rdv3z+zGPJvjhw5AiHxQiNu9PL5+OOP7VW6dA1f2bJlhp6s5QAAEABJREFU7eQjJV7qRRASLyzh7rBlHM5wAqgQe+mll7jKl+FxpnskQAIkQAIkQAIkQALOEnAl8TN9+nRbGmSaxGvUqFFy3XXXORsBhUbDFXVUYbRu3VqOHj2qsaTiH6hYxStt2rT/POC/yhCYMGGCfPjhh6JrFZkyIANoCI6ZTp06yf79+7X/fBL+RwIkQAIkQAIkQAIkQAIeEXA88fPll1/aVRZK9O1wEGLTpk2lbt26YlmWg6OqM9T//d//Sa1atWTs2LFGXE1HX58KFSpQ4qXOIWZb8t1330nPnj1t+aD9BP8hgQgJIEFNyVeE0Lg5CZAACZAACZAACQSBAH1MlIDjiR9IvA4ePJjohDq+cMMNN8iIESPE1FW80M/nsccek3fffVdraVfo2LrvvvsEiR9KvEJE1LmHxOvkyZPqGERLtCSwZs0amTFjBhOIWkaPRpMACZAACZCA+wQ4AwmQwOUEUlz+MLZHkyZNktdff924EnxIvLJmzRobHEX3XrdunV3JhIofyCgUNTNss9B0u0+fPkKJV9jIPNtw5MiRsmnTJkq8PCNu7kT4rIIs9fvvv+fxZG6Y6RkJkAAJOEGAY5AACZAACfxNwLHEz+effy5xcXFimsQLkoLatWsbKfFC0qdKlSqCCi1T+jEh6VOuXDlKvP5+c6v0/+7duwXVgGfOnFHJLNqiMYGQ5IsVZBoHkaaTgKcEOBkJXE4gTZo0Rv6+v9xLPiIBEiCBfwg4lvjBSd2RI0f+GdWQf/PkyWOsxCuU9Dl16pQh0RKpWLGidO/eXSjxUi+kjRo1Ep6gqxcX3S2CTHXq1KmUfOkeSK/tN2w+JEENc4nukIAnBNKnT8/EjyekOQkJkIAKBBxJ/IwbN06WLFlinMRr9OjRkjlzZhXi5KgN3377raDSx6SkD67aoNqHEi9HDxVHBhsyZIhs27aNkhxHaHKQSwlA8tWjRw/Zu3dvVMfXpWPx738IIHFuWWYuYvCPh+b9e/r0aeN+f5kXJXqkIoF06dIx8aNiYGgTCZCAKwRiTvxs3bpV+vXrZ5zEq3Xr1lKjRg3jvhBwolSsWDExKemDdwaSPnfffbdx8YJvOt927NghSPycPXtWZTdom8YEUO3QrFkzVpQ5FENeAXcIpIfDMPHjIexkpsLnUTKb8GWFCODzDr0hFTKJppAACZCAawRiTvygr8/vv//umoF+DJwvXz5jJV4NGjSQP/74ww+srs2J6qXOnTs7IPFyzcTADtywYUOekAc2+t45vnHjRpk4cSIlXw4gx4mQA8NwCA8J4EKOKX36PMTmylRMwrmC1bVB8XlnWZZr43NgEiABElCJQEyJH6zS884774hpPzhGjx4tGTJkUClOjtgyfPhwe9U1VP04MqACg6BMF9U+lHgpEIx4JqASEKvFmfb5EM9NPlSAAD7T8DmAJuL4WwGTtDUB332WxRMhnQKIZINO9ppsK2LB7zx9IozfkJbFzzt9IkZLScBFAgEYOurED66w9u3b17i+Cu3atZOqVasaJxl65ZVXbEneuXPnjDqscQyWLl3auHjpHqQtW7bIqFGjxLTjTfe4mGw/JBaQfJ04ccJkN133jVfAXUfs+ASs+HEcadQDMvETNTpfduTnnS/YlZ+UBpKAqQSiTvxA4oUfGyaBKViwoJESr88//1y6du1qnAyievXq0rFjR0mVKpVJh6ERvnAVLyPCqJ0Tn3zyiYwfP964zzovA8ETIS9pOzMXkw3OcHRiFMbCCYrejZHM5513hnAmEiABEvCAQFSJn6FDh8rKlSs9MM/bKbCKF1aH8nZW92cbNGiQ/Pzzz+5P5OEMmTJlkt69ewslXh5CD3MqrLK0Z88e4ySgYbrPzXwkAIkFLkpAYkjJV3SBwIlQdHtyL78I/Pbbb8ZVX/vFMuF5w38WF0TxORT+HtzSTwKUevlJn3OTAAl4TSDixM+6desE8hrTflR36tRJKleubJxkaMKECbJ48WLjfhTiGCxZsqRx8fL6A8Dp+TZs2GBXXFDi5TRZjhcuAUi+mjdvLpR8hUvs8u3Y4+dyHpc9UvTBjz/+aNx3vKKokzWLFT/JIlJqAyS6uaqXUiGhMSRAAi4SiDjxg6up+GHtok2eD3377bfLsGHDJGXKlJ7P7eaEW7dulQEDBhjXZ+Xhhx+WZ599lhIvNw+eKMeGxAtXPKPcnbuRgCME8NmHCk6chDkyYLxBTH6IEyHLskx20TjfDhw4IBcuXDDOLx0dYvWVXlFjxY9e8aK1JEACsRGIKPHTv39/WbNmTWwzKrj3mDFjJHXq1ApaFptJgwcPll9//TW2QRTbO0uWLILVeyjxUiwwf5uDfks//PBDkCRef3vN/1UkAKkFJK47duzgCXGEAULiJ8JduLnPBFDxw8SPz0G4OD2TcBdBaHKHzzvLYqJbk3DRTBIggRgJhJ34+eCDD2TgwIHGlRN369ZNypcvHyNG9XYfO3asLF261Lh4QeJVtGhRxSRe6sXfa4tWr14tL7zwgphWDeg1R87nHAEci5B8nTx50rlBAzAST4T0CzJ6+J09e1Y/ww2zGNWuR44cMe53l2Fhuswdft5dhoMPSIAEDCcQVuIH/Tri4uKMu3JavHhxO5nlmMRLkYMFK9ugOgtxU8QkR8x49NFHpXXr1pR4OULTuUHQ76thw4bCE2znmHIkZwhgRcMRI0YIJV/h82SPn/BZqbIlKtxYaeJ/NH766Sfjfif7T9VdCyj1cpcvRyeBQBDQyMmwEj/oE4OmrRr5FZap6AFhosQLEodjx46FxUCXjbJly2ZLvExcdU2XGCRmJ/otHTp0KLGX+TwJ+EYAJ8RYhXLbtm08IQszCrwCHiYoxTaj3Mv/gDAG/scgUgvwecfmzpFS4/aJEeDzJKA6gWQTP8uXLxf8cMZVfdWdicS+Xr16SdmyZSPZRYttR44cKcuWLTOuzwokXoULF6bES7Gj8L333pNZs2ZR4qVYXGjOvwTQ+4SSr395JPcXToLQQ82y2PciOVYqvb57924mN30OCKuufA5AFNMj8WNZjn/WRWEJdyEBEiAB9wkkmfg5ceKEmCjxKlWqlPTr18+4Vbw2btwokHiht4X7h453Mzz22GPSokUL4+LlHUF3ZoJ8pkGDBpR4uYOXozpI4P/+7//sCxjoweHgsMYOhZMhY50z1LE9e/YwAe9zbP+b+PHZIE6fLAFKvZJFxA1IgAQMIpBk4gcSr82bNxvk7j+umCzxQrLuHy/N+Pfaa68VVPtQ4qVePNu2bSumSQrVo0yLnCAAyRd6/WzZsoVVEWEAzZgxI6srw+AU9iYebPj1118z8eMB56Sm+Oqrr8S03opJ+WvCa0hyWxYrfkyIJX0gARJInkCiiZ+3335bRo0aZdzqBEgilClTJnkymm0xfPhwgSwPJziamZ6kuajMKlCgAE9CkqTk/YtvvfWWzJs3jyca3qPnjFESCEm+/EyOR2m657vxZMhz5DFPSKlXzAhjHmDXrl1M/MRM0dsBWPHjLW/ORgIk4C+BBBM/R48eNVLidffdd0vv3r0FPQz8xe7s7B999JGgOss0idcTTzwhTz31FCVezh4uMY/2+++/S6NGjYSymahRckefCODkePDgwVzlKxn+TPwkA0jBl7GiFD6bTbv4oyDqBE1CpQ/ldgmiUfrJTJkyGXdOoDRwGkcCJOArgQQTP0gibN261VfD3Jh8zJgxYuoqXqadhOfMmTMAq3i5cZS7PyYkXsePH3d/Is5AAg4TwEkxpL7oh4YKIIeHN2Y4JH6McSYgjuDYptzLv2Cz2sc/9rHMfPXVVzPxEwtA7ksCJKAVgf8kfhYvXizjx48X/IjQypNkjB04cKCULFkyma18ejmGabHi2sqVK42LFyRe+fLlo8QrhmPDjV0XLlwor7/+OvukuAGXY3pCAJWRzZo1E0q+EsedOXNmfvYmjkfZV9CTEZUnyhposGFM/OgZXCR+UqZMqafxtJoESEBvAj5Yf1ni5/Dhw/ZqV6ZdCS1Xrpx069bNuKz+unXrjFzFq2HDhtK4cWNKvHz4QEhqSnw+NG/enBKvpCDxNS0I7N271/7sNK1S0in42bJlM+770ik2Ko+zfv16OXv2rMomGmvbjh072N9Hs+imSpVKkOQ2rf2DZmGguSJCCCTgFYHLEj/9+/cXfHl5NblX85go8frzzz8FVUym/cjLnTs3V/Hy6o0R4Txt2rRhlUSEzLi5mgRQ0YrK1g0bNrBBeQIhYuInASgaPIWLQX/88YdxFcAaoBd8lpw5c0YHU2njRQJZs2ZV7QLjRct4RwIkQALuEPhf4ue1116TKVOmGPeDYciQIXLHHXe4Q8/HUeHX6tWrjYsXJF558uShzMDHYyuhqefMmSNYycu0asCEfOVzwSAAyRcq2E6ePBkMhyPwEokfy+ISxxEgU2JTVLBhsQfKvbwNx/79+2Xnzp0OVvx4a39QZ2PiJ6iRp98kEFwCduIHq0H07dvXuL4dFSpUkM6dOxtXsr5mzRoZNGiQcVeqmzRpIg0aNOAVGFHrvx9//FFQ7XP69Gm1DKM1JBAjge+//95e6REnzDEOZdTuSPxQ/qBASKMw4eOPP6bcKwpuseyCZvH8foyFoD/7or8PP+f8Yc9ZSYAE/CFgJ34g8UJjOn9McG9WEyVeuJJnosTr5ptvtvtLmbjqmntHuDcjt27dmhIvb1BzFo8JQPI1efJkWbt2rdKJdI+xSPbs2Y27YOI1Q7/mg9zLNAm4XyzDnRfJNsq8wqWlznas+FEnFrSEBEjAGwIp5s6dKzNmzDBOMjRixAgpXLiwNxQ9nAUSrw8//NC4eEHidcMNN1Di5eGxFM5UL774oqxYscK4asBwfFdwG5rkAgHIFyH54ipf/8JFxY9lUer1LxF9/tq2bZv9mc3kj3cxQ+KHvL3j7dRMTPw4RZLjkAAJ6EIgxTfffCM9evSwy9179+5tzH3Hjh2Nu2L5/vvvC5ZvR28KXQ6wcOzE0sqPP/44JV7hwLK38eafb7/9VvA+Ygm7N7w5i38EIGfs2bMnV6y7GAIkfiiBuAhDsztUsS1dulT4ue1N4MB69+7drBj0Brejs7Cy0VGcHIwESEADAinwYxfVFqbdsEyjBvzDNhE/4tDX57KrSmHvre6G+fLlsyVeadKkUdfIgFpGiVdAAx9At3GyPG3aNPnggw94Avd3/HFCZFnW33/xfx0JoBH/nj17WKnpQfCWLFkibBDvAWgXprjxxhuF7QVcAMshSYAE3CUQw+gpkCAx8RYDEyV3hcRrw4YNxkm84uLiJEeOHEoyD7JRWOEPq8b9+eefQcZA3wNEAJIvVB/+/vvvAfI6YVdxMoTPZVb9JMxH9WePHz8uqERh3xl3I/XZZ5/J22+/zUpBdzG7NjoSPzj/cW0CDkwCLhPg8CQQKQG7uXOkO3F7bwksX75c0LPINIlXq1atpE6dOpR4eXs4JTsbGr1369ZNeI6Sxe0AABAASURBVNKQLCpuYBiBQ4cO2dJnrvIlgob7KVOmNCzCwXFn8eLF8scffxh3sUilCKKyir3BVIpIZLbcdNNNYkjiJzLHuTUJkEBgCTDxo3jo8aPCRIlXwYIFpW/fvkKJl3oH4DPPPMNVvNQLCy3ygAAkX2hojmS7aYn2SPEh8cOKn0ipqbP99u3b5YUXXmA1ikshOXjwoLzxxhuK8XXJWQOHRR+zK664wrheoAaGii6RAAk4SICJHwdhujHU4MGDBStG4ITEjfH9GhMSr2uvvdav6TlvIgTGjh0rH330kVDilQggPm08AUi+sMrXr7/+aryvSTmIxA8rfpIipPBrf5uG3wzTp08XLOCBY/rvp/i/gwRmzpwpWACBbB2E6uFQuXPnZn8fD3lzKhIgATUIMPGjRhwStOLdd9+VUaNGGdegsW3btlKrVi1eaUkw6v49uWPHDunTpw8lXv6FgDMrQuDw4cO25Ev3pq2x4MyTJw9luLEAVGDf/fv3C5I/QT6O3QjDzp07ZerUqYJeSm6MzzHdJ0CZl/uMOQMJkIB6BJj4US8mtkXQ5g8cOFDOnTtnPzbln8KFC1PipWgw0XOJJwiKBic2s7h3FARmz54ty5YtM+4zOFwUSPxQ6hUuLTW3C1X9bN68mavVORgiVPscOXKE/ZMcZOr1UKz48Zo45yMBElCBABM/KkQhARvQ1wc/1vDDLYGXtX0KEq+sWbNqa7/ehiduPZqHb9myhRKvxBHxlYARgISjRYsWguof0z6Hwwll3rx5WfETDijFtzl9+rSMHz+efdscitPatWtlzpw5woskDgH1aRhW/PgEntOSAAn4SoCJH1/xJzw5lgdFrxWceCS8RYzP+rR7hw4dpGbNmpR4+cQ/sWk//fRTQXXZ2bNnE9uEz5NAIAn89ttvtuQriKt8pU+fXnLkyMHPawOOfPymmDFjBhsRxxhL/CabNGmSHDt2LMaRuLvfBJDYTp06td9mcH4SIAES8I7A3zMx8fM3BJX+P3r0qH0SbprE64477qDES6UD7RJbWrZsyavBl/DgnyRwKQFc3V+6dGkgJV9s8HzpkaDv32jWP2bMGEEfNyQv9PXEX8txQW7FihXsg+dvGGKeHU3r0XaAq8rGjJIDaEqAZgeXABM/isUeEi9UYJgmLejXr59kzpxZMdo0Z8CAAfbJgGnHGyNLAk4RwIkyVvk6ePCgBO19UqJECUmVKpVTKDmOjwQOHDggSFycOHHCRyv0nXrNmjXy/PPPs9pH3xD+z/IiRYpIunTp/vc4wH/QdRIggYARYOJHoYAvWbJEJkyYYNwqXs8995xUr16dkgGFjjWY8tFHHwl6+1DiBRq8kUDiBNBsv0ePHoGTyhQvXpyJn8QPC61eQdJy4cKFgmQ/V6OKLHR4/+O7EtLPyPbUZetg2YkKdFb7BCvm9JYESOAfAin+ueO/fhPAChH4QWaaxAtXjFHtQy2130fYf+eHxIsNKv/Lhc+QQEIEXn31VXnzzTcDJfnC5zcrfhI6GvR8DtVrqPrp0qWLIJnxHy/4RIIEunfvLh9++KHwIkmCeLR7smjRopI2bVrt7KbBJEACJBArASZ+YiXo0P6QeG3fvt04KUFcXJxkzJjRIUocxikCvXr1kl27dhl3vDnFh+OQQHwC6JMCydf+/fuNf9+EfEdz5xtvvJHVmiEgBtwj+TN9+nRh8ie8YLZt21Zefvll9sELD5cWW6Hih4kfLUJFI0mABBwmwMSPw0CjGW7RokUyefJk4yReXbt2lSpVqvCkIZqDwsV90Ktg7NixgapccBGnyUPTt3gEsLpXkCRflmUJ+mGw6ifegaD5QyR/sMrXkCFDmNBIIpZI+syePVsojUsCkmYvZcmSRfLkyUMJq/A/EiCBIBJg4sfnqB86dMjW3Jsm8SpdurSg2ocSL58PsHjTo1QdVQs4gY33UhIP+RIJkECIwGuvvSa4mfaZHfIv/j3lXvGJmPEYyR80K0YPPsq+Lo/pvn37BN+Ts2bNYtLncjTaP4LMi42dtQ8jHSABEoiSABM/UYJzarfBgwfLF198oYd0IAKnkfThl2sEwDzaFBKv7777zrjjzSN8nIYEJCT5+vbbb+2/TUfCxI+5EUbyB5U/zZo1EyQ70ADaXG/D8+z999+XevXqySuvvMJqqPCQabUVZF5s7KxVyGgsCZCAgwSiSvw4OH+gh8JV42nTphkn8YIUomLFimJZVqDjq5rzK1assCWF58+fV8002kMCWhFAtQ+SqEGonEPiJ0OGDFrFh8aGTwDJnzfeeMNOdqxcuTLQEuBJkybJ448/Lps2bZIzZ86ED5FbakOgZMmSwouS2oSLhipGgOboT4CJH59ieODAASMlXvfcc4/06dNHKPHy6cBKZFr0KEDpehBOVBNBwKdJwFECWOFrwYIFxq/0g5Mk9vlx9NBRbjAkf5DsqFu3rowZM0aCttojJPfPPvus9OzZU7DCKqr6lAsSDXKEwJ133skVvWInyRFIgAQ0JcDEj0+BQ1PFL7/80jjJDSVePh1QyUyL6gQkG1nKnwwovkwCYRLAyXKrVq1kz549xku+0BeDDZ7DPDA03QzJjt9//1169+4tOK737t1r/HGNUGGZdki7IHmD/3iOt3AJ6LUdZF6ZM2dmNbpeYaO1JEACDhJg4sdBmOEO9eqrrwqWU8WJQ7j76LAdKn3Kly/PL1XFgvXOO+8IftRS4qVYYGiO9gSCIvli4kf7QzVsB3BMz5s3T+rUqSMzZ86U06dPJ7+vhlvAT/wOe/TRR2XdunXG+qlhaFwzuXjx4qz2cY0uByYBEtCBABM/HkcJDRQHDhxonDygXLlygt4+vCrs8QGVzHS//vqrPP3004Er3U8GC18mAccIvPXWWzJ37lzjPtMvBYQTpnA+2y/dh3/rSwAXpT777DNp3bq1IDHy9ttvG3N8//jjjzJ+/HjBb5b27dvL4cOHA1HZpO/R6Jzl6O+TNm1a5wbkSCRAAiSgGYEUmtmrvblYxWvXrl3a+xHfgf79+7NhXnwoCjwO9SxQwBSaEBwCgfIUEpm2bdsKPtfxtxj4X4ECBeTGG2+UFCn4k8HA8CboEmTBqIp577337ORPy5YtZfny5domgJDIguS5bNmy0qVLF9m8ebOw512CoTf2Sfb3MTa0dIwESCBMAvwVFyYoJzabM2eOvPTSS8ZdXULSBz+mLIureDlxnDg1xqJFi+Tll18W/yReTnnCcUhAbQI4QUaS1dQTScuyhHIvtY9Bt6xDMvPs2bPyyiuvyEMPPSRYsfP555+3e1uhMsiteZ0aFyuVoeoVMvQRI0bI999/byev4JdTc3Ac9Qmgv0+uXLkkZcqU6htLC0mABEjAJQJM/LgENv6w+LFhosTr/vvvt6+eJSkDiA+Dj10ncPDgQVviZeqJqOsAOQEJREhg2bJlMnv2bPukMsJdtdiciR8twuSakUjyIAH00UcfCSpnSpcuLY0aNbIvLuD7RpVECuxENc+UKVPkwQcflEceecR+Xx47dsy+CIJKJtcgcWBlCUCuihUKlTWQhgWCAD6XoPzgbbCQQewMPv/8c8GFx0jePCki2TjWbYO8/6BBg+wrZKYxGDBgACVeCga1e/fu8scffyhoGU0iATMJ4MS3Q4cOsnPnTuOqOhExXDFngh8kgn3DcY4fmugft2DBAmnWrJkUK1ZMGjZsaCeBIHlE8sUrSri4sWbNGsFKqUjy5M2bV1Ddg/49kKYdP35cvLTHK785T2QE7rrrLjZ2jgwZt3aYwOTJkyUuLi6Qtzj67XjcLcuSG264QVKnTh3RkcrET0S4otsYchvc8IMpuhHU3AvZWnyZWpalpoEBtWr+/PmycOFC/tgNaPzptn8EcEKMaoiTJ0/6Z4RLM+NkOl++fOzz4xJfHYdFQgXH/KFDhySUBEJl2G233WavCjZ06FBZtWqV7N+/P+ZkKN5TGGf79u2C6rq+ffvasrPrrrtOqlatav+oRqP1H374we7dA7tM+82l4zGigs1Zs2a1m3mz4sf3aATaALRd4O28XXlJDrFxQDUrLrZcddVVEb+nmPiJGFlkO+zdu1dQFYMS6cj2VHvrypUrS6dOnYRXgNWKE370tmjRwv7hq5ZltIYEgkEAVQYzZ840TvKFxs7o75ImTZpgBJJehk0A8qlQEujMmTOye/duWbx4sSA5gx+oSBhee+21ctddd0mDBg3s30S4cBT/1qdPH0Gj9CeeeEKQzClVqpTccsstgh+32bJlE4wDiVnt2rUFSSVU+vz++++COZHogQ2wJWzDuaFPBLydtkyZMoLkDz7DvJ2Zs5EACZCAswTQqwzV5bjoEc3ITPxEQy2CfdDX59tvv41gD/U3xZcnkllcFlO9WHXr1k1wZVQ9y2gRCQSDAKoMOnfuLNBe40TUJK+R+Im0rNgk/+lLeASQfMGxj6u6uOh1+vRpe9n0Tz75RFCRit9Fdul/vPL/YcOGybRp0+yKVTRl/vTTTwUXz44ePWpfzMA4SPJgTIyN91p4FnGrIBO455572JIgyAcAfScBgwi0atVKIL2P9rcYEz8uHgxYwWvu3Lkxlzi7aGJUQ+NKW8mSJcWyrKj2507uEICcEFdZ8YPbnRk4KgmQQDgEUH1g4ipflSpVkltvvdXxlXHCYcpt9CeARA2+n5C0SeqGbbAtbkgi6e85PfCTABM/ftLn3CRAAk4RqFWrljRu3Niugo12TCZ+oiWXzH5ff/219O/fP+Ju28kM6/vL1apVEzRNpMTL91BcZsCePXvsEnlcEb3sBT4gAX0IGGUpeptMnz7dlqGY4hiqPVn1Y0o06QcJmE8AMi8kq/mb1fxY00MSMJkAGjnj/DtaiVeIDRM/IRIO36OUed++fQ6P6u9w6O2AZBbu/bWEs8cnAInXiRMn4j+t6WOaTQL6E0C1At6X27ZtM6rROhI//A7Q//ikByQQBAJ33303ZV5BCDR9JAHDCcQq8QrhYeInRMLB+xkzZtg6dvzwd3BY34eCxAvLtlqWRxIv3z3WwwAcb1jlxLTjTQ/6zlrZpUsX4Um1s0z9HC0k+TKp71ZI7oXqHz/Zcm4SIAESSI5A2bJlmfhJDhJfJwESUJoAFjTAKl5ZsmSJ2U4tEj8xe+nhAF999ZWREq+aNWtKmzZtuIqXh8dSOFPt3LnTXl2NEq9waKm9TdGiRaVXr16C/lk8qVY7VpFYt3btWpk6daoxki8cm0j+MEEZyVHAbUmABLwmgGqfu+66S7gQidfkOR8JOEOAo4jceOON0rFjR4lV4hViycRPiIRD94MGDZIDBw44NJoaw6RLl85OZvGHvhrxuNQKSEko8bqUiL5/d+3aVTJmzCi45w9VfeMY33JU4qHR8+bNm42RfFWuXJmVafEDzcckQAJKEahatar9naqUUTQmGgLchwQCS+CZZ56R22+/XaJdxSs+OCZ+4hOJ4TGu6i5cuNC4VbyGDx8uRYoDfVQkAAAQAElEQVQU4SpeMRwbbuw6efJkef/994073txgpfqYjRo1klq1atkVdTVq1BCUdDr1Ia+670GwDysYoZrLFMlX+fLl5bbbbuPqXkE4eOkjCShDIHxDrrzySsFiJBkyZAh/J25JAiRAAgoRePTRR+XJJ58UJyReIbeY+AmRiPH+iy++kAEDBhi3itfDDz8sLVu25A/8GI8Pp3f/7LPPBNU+Z86ccXpojucxAZRv9ujRQ9KnT2/PnDJlSsHjHDlyMNlqEzHjnw0bNsjEiRPFBFkm5F6s+jHjuNTSCxpNAskQeOCBByRPnjz87ZoMJ75MAiSgJoGbbrrJUYlXyEsmfkIkYrwfPHiwHDp0KMZR1No9U6ZMgtXJKPFSKy6wBk2ATakegD9BvkHadfPNN1+W5IGmF4k9vvfMOTIg+erbt698/PHHggog3T1TIfGjO0PaTwIk4A6BKlWqCKt93GHLUUmABNwn0Lp1aylcuLCtBHByNiZ+HKA5adIkef31142T3AwbNkwKFSp02QmpA7g4RIwExo4dK+vWrTPueIsRi5a748dps2bN/tMvxbIsadGiheDkGhVAWjrnndHazISET+/eveXUqVPa2JyYoWicWrx4ccd/lCQ2H58nARIggXAI5M+f35Z5hapow9mH25AACZCAKgTq1KkjTzzxhGTOnNlxk5j4iRHp9u3bbYkXftDHOJRSu0NXiBNSnnQqFRbZtGmT9OnTx5gVgpylq99oqOpJ7KpkqlSpbDkfGz3rF9ekLN64caOMGTNGe8kX5V5JRZmvkQAJ+EUAF1RQse7X/JyXBEiABKIlAAVAp06d5Nprr412iCT3Y+InSTzJvwiJ1y+//JL8hhptkSVLFr0lXhqxjtTUzp07CyVekVJTc3tIvMqUKZNkDwJUVXTo0OE/FUFqekSrwiEAyRcktKja0/2CAU6wKEcMJ+rchgRIwAsCSPjUrVuXq3l5AZtzkAAJOE4AEi+obXDx1/HB/x7Q6MTP3/65+v/48ePlzTfflL/++svVebweHBKvW2+9lRIvr8EnMx9WV9u8eTMlXslw0uHlO+64Q3r27Cnp0qVL0lxUVXTv3t3W+eLvJDfmi9oQQMIHki/dk7g4juvVqyesStPm0KOhJGA0gdq1a0vBggUpQTU6ynSOBJImoOur+D31+OOPuyLxCjFh4idEIsL7rVu3GinxwkHXpEmTJKsQIkTFzR0gsH79ersK6+zZsw6MxiH8JoBqn4wZM4ZlxhVXXCHYnifXYeHSZqMtW7bIqFGjtJZ8IRmJpUaTS2BqExQaSgIkoDUBtCnAd6bWTtB4pwhwHBLQhsAtt9xir+LllsQrBIKJnxCJCO8HDRokv/76a4R7qb15tmzZ7OQCS/fVitOFCxfkueeeo8RLrbBEbU3jxo3loYceiuiKJBq9PfbYY5I6deqo5+WOahGA5Gvo0KGyZs0arVf5uueee+zjmd8bah1ftIYEgkagRo0aAvn0fz+LgkaC/pIACehG4JlnnpECBQpEdG4QjY9M/ERBDasqLV261DiJF05C0FTKsqwoqHAXtwggLp999plxx5tbvFQe97rrrpMePXpIpKuNoMk6GkFfeeWVKrtH2yIkcO7cOdFd8oWqn/r16ycrW4wQDTcnAecJcESjCbDax+jw0jkSMJYAfkNBcePFb3wmfiI8jFCe379/f62v0CbkMpaNa9iwISVeCcHx8blVq1YJei5R4uVjEBycGsmbPHnyCE6WIx0WVwLQF4iSr0jJqb39tm3bZMSIEVov8f7AAw9I5cqVtalIU/uIoHUkQAKREihbtqxUrVo14osqkc7D7UmABEjASQKQeIVW8bIs9wsvmPiJMHpYjeXYsWMR7qX25tATQrrGE0q14nTq1ClKvNQKSUzW4Efp008/HfUKXZZlSbt27eTee+9lgjamSPxvZyX+gOQLjduR5EXTZyWMitAIVKShISF7/UQIjpuTAAk4QuCRRx4R9vZxBCUHIQES8JBAmzZtBAsqubWKV3xXmPiJTySJx2jEuWzZMuMkN6goyZ07dxKe8yU/CAwZMkT+7//+z7jjzQ+WSc/pzauo9smQIUNMk+GLAeMwSRsTRuV2RsIHkq/jx48rZ1u4BmE1HfT7wTEa7j7cjgRIgARiJVC4cGFBD7xwF0yIdT7uTwIkQAJOEMDiGPjs8jJpzcRPmJHbtGmTxMXFGSfxgrwLMi9csQ0ThbmbKeTZu+++K2PGjBH0AFHILJoSJQEka0qXLu1IpU6FChWkVatWUVcORekCd3OZwI4dO2xZJyr9XJ7KleHxHQKdOqt+XMHLQUmABBIhgKSzF70xEpmeT5MACZBAxATy5csnXkq8QgYy8RMiccl9Qn9C4nXixImEXtL2uZw5c8rgwYN5AqlYBI8ePWp/GJw8eVIxy2hONASKFSsWVUPnxOZCfyAkkvLkySOW5b4eODE7+LyzBCD5QlXpihUrtE344iJC8eLFXV+VwlnyHI0ESEBXAtdff72gKSqrfXSNIO0mAXUIeGlJ27ZtBf19cNHMy3mZ+AmDNhpvLl++3DjJDfpK5MqVKwwC3MRLAljF65tvvjHuePOSoUpzde3aVZz+UZo9e3Y7mUTJl0qRjt0WSL569eolSP7GPpr3I0DmhV4/PC69Z88ZSSCIBNDbB79jcUEkiP7TZ1cIcFAScJVAgwYNpE6dOr70JWPiJ5nQfvTRR0ZKvJo2bWprovllmcwB4PHLS5YskYkTJ2p7xd9jXMpP16RJE6lZs6bjFRCWZQm+OB5++GHHx1YequEGoq8XkvK6Sr5Q9ZM/f35HZI2Gh5rukQAJxEAgU6ZM9skT7mMYJold+RIJkAAJOEsAjZz9kHiFvGDiJ0QikXtIvE6fPp3Iq3o+jUbOlHipF7tffvlFOnToILqe8KlH1F+LsFpez549XVteFuWhkHyxusLfODs9+19//SXjxo2Td955R8sEcObMmQW9fnhcOn1kcDzfCHBiJQmgt0+hQoV48UPJ6NAoEiCBhAhgFa+8efP6dnGMiZ+EonLxOUhu3n//feMkN7iajJPSi27yThECWMVr//79xh1viuD13Izu3bvLjTfeKG5W1RUtWpSSL88j6/6EkHxhla/Dhw9r+XmAqp+bb77Z1WPf/Sj8dwY+QwIkoAaB9OnT2719vFwNRw3PaQUJkICuBBo1amRXKfr5ucXETyJHz7p166R///7GreLVvHlzgSbazZPRRJDy6SQILFy4UF544QXjjrckXDb6pWrVqsnTTz/teuN0vI9R9VOyZEmeZHt3RHky0+7duwX95XSsOM2RI4egcSFOzjyBxUlIgAQCRaBFixZy9913u/4dGyiodJYESMA1ApDA+ynxCjnGxE+IxCX3WF0FEq+zZ89e8qz+f+IK7KBBg/hFqVgoUeUDiRdX8VIsMMmak/gGqPbJkCFD4hs4+Aoa6mI+SmschKrAUJB8od/XW2+9pZ3ky7IsO/FZpUoVSZ06tQI0aQIJkIApBFDpit9MWbJkMcUl+kECJGA4gWeffVZwHo42DX66ysRPAvQhuVm9erWWJfYJuPO/pyDxypYt2/8e8w+HCMQ4DI63n3/+OcZRuLsqBJCEKVWqlKf6XVQYoWE7T7JVOQqcsQOSrz59+sihQ4e0+z5CQpJVP84cBxyFBEjgXwKo9rnmmmvEsqx/n+RfJEACJKAogcaNG9tqGxUa0TPxE+8gQcIHVTH4wR3vpWQfqrxBq1at5KGHHqIcRLEgzZ07V2bNmiUXLlxQzDKaEw2BO+64Q5D4SZcuXTS7R70PriBgXv4Yjhqhsjvu2bNHhg0bpmXT9/vuu0/w3eP1+0HZYNIwEiCBmAhgCWTcvKqojclY7kwCJBAIAkk5WaBAAencubOgt65l+Z+sZuLnkmidO3dOkPQxTeKVL18+4SpelwRakT+//fZbadeunZYndIogVM6MHj16SMaMGX2x6/rrrxc0BKbkyxf8rk0Kyde0adPkzTff1E7yhR5U+IzDDx/87RokDkwCJGA8gaxZswpWxLnqqquM95UOakmARpPAfwhA4uX2Qi//mTSJJ5j4uQQOJDcffvihdiX1l7iQ4J+QeFELnSAaX59EkvHYsWO+2sDJnSMAqVX16tV9W1rWsixp1qyZPPDAA77Z4BxNjnQpAVSgQvL1448/avf9lDNnTjZ6vjSY/JsESCAqApB4FStWTIO+YVG5x51IgAQMI4Dzgtq1a4sKEq8QWiZ+LpJ4//33ZejQocatqoSrIzgZ5dXWi4FW5G7mzJkyb948SrwUiUesZmTPnl169uwpfq9ihL4q3bt3ZwP3WAOq4P7fffed6Cj5sixLoG/H91CaNGkUJEuTSMAlAhzWMQJlypSxE8iZM2d2bEwORAIkQAJuEShYsKBSEq+Qn0z8/E0Cy+WauIoXyutRxcQf238HWaH/v/rqK2nfvr3guFPILJoSAwFIrHLnzq1ED63SpUsLlozk+z6GgCq4KyRfL774orz22muimxwZCckgN3pW8HCiSSSgFYGWLVsKJV5ahYzGkkCgCUDmrsp5waWBYOLnbxpIjmzYsEG7Evq/TU/y/xEjRihVXpaksQF6Ef2WTp48GSCPzXY1tKKWKokWVPeh1xCuNuBvs+lr5V3MxoYkXz/88IP8+eefMY/n5QD33HOP3Z/D76o4L33mXCRAArETaNSokaBikJ8dsbPkCCRAAu4TePrpp6VWrVq+9fxMysPAJ36WL18uSJCYtqpShw4dpGrVqkpUICR1AAbttalTp8rrr7+u3Ulb0OIUib9IskS2wkgko0e3Lezp1auXsNFzdPxU3mv//v225Eu3ikEkIdHksFChQoJV6FRmTNtIgATUIIAr5vjcYLWPGvGgFSRAAkkTwG8crOKl6iq7gU78oOrCRIlX4cKF7VW8UqdOnfTRyVc9JbB9+3bp0qXLvxIvT2fnZG4QQF+fO++8U8kT2UceeUTq1avHRphuBN7HMSH5mj17tsyfP187yReWM6Xky8eDh1OTgGYE0NA5f/78Sn7HaoaS5pIACXhAAIUXWGUXF7s8mC7iKQKd+IHkZuPGjb5LvCKOWjI7oIKJJbHJQPLhZRxvSDb6MDWndIFA0aJFpVu3bpIuXToXRo99SFRVqFiNFLtnHAGSL/SV2rt3r1bVg5ZlScOGDe0SaFWkkTyaSIAE1CSAFSohmbjiiivUNJBWkQAJkMAlBLCybs2aNcOWeF2yq2d/Bjbx8+6778qoUaOMW1Xpueeek0qVKollWZ4dRJwoeQLjx4+Xt956S6uTtOS9CvYWkFJlzJhRaQj58uWTfv36UfKldJSiM+7gwYMyfPhwOXXqVHQD+LQXEpKDBg2SEiVK8Cq+TzHgtCSgOoHbb7/dlrRixUzVbaV9JBAjAe5uAIHbbrtNyVW84qMNZOLnjz/+kAEDBmhXJh8/ePEf33HHHYIf1JR4xSfj7+MtW7YIkgRnzpzxBjWXeQAAEABJREFU1xDO7hiBp556ym42iZNYxwZ1YSDLsgT9EdBYV3VbXXDf6CEh+XrllVdk7ty52n2XoW8HKiBVT5wafQDRORJQlAB606ENA3plBOt7S9GA0CwSIIFkCXTo0EFy5cqlfOFFIBM/SI588skngh/OyUZSow0g8cIXpkYmB8JUJBkp8TIn1NmyZRP09tFFTpkqVSqBLIifDeYcgyFPsChBnz59ZNeuXVpVE1qWJffdd5+MHDlS6ZLoEGfek4CnBAI+2ejRo6VixYqsVA34cUD3SUAXAi1btpQaNWpo8XsmcImft99+W8aOHWucxAu9Ru6//37lM426vImdshMnNitWrNDqpMwp300dp2/fvoKKBcvSR06Jk+zWrVsL+6qYd1T+8ssv9sqUukm+0PiwadOm0q5dO9Elier10cP5SCBoBFAd/cQTT0imTJmC5jr9JQES0JAAFlTCKl5YvMKy1D8vCFTi5+jRo0ZKvNArAX08cGVfw/eMsSZ/9NFH0r9/f6HEy5wQV6lSRXCyqpucEifZ3bt3lxtuuIHJYf0OxyQtRuUqVvh6+eWXtfuswXcWPiMhRcTfSTrKF0mABIwm0KRJE2nfvr1kyZLFaD/pHAmQgDkEOnToIDly5NDmt3WgEj/oKbB161ZzjraLnjz//PPKrix00cTA3eFkjBIv88IOyZR/1Qmx8bz66qsFCWJKvmLjqOLeoVW+vvzyS+2qWZFEnTJlinDlHhWPLNpEAt4QKFKkiH1hFlJqb2bkLCRAAiQQGwFIvKpXr66FxCvkaWASP0uWLBGsrISeCCHnTbjHcs333nuvNplGE5iH48PQoUNl9erV7ki8wjGA2zhOAH19SpYsqe1KRJZlSf369aVmzZrC6grHDw/fB/z111+1lHwBXN68eeXVV1+VzJkz4yFvJEACASMwffp0yZkzJ3/LBizudJcEdCWAlQexkjYkXjr5EIjEz5EjR+wrCWfPntUpNsnaWrp0aUFjT57EJYvK0w3WrFkjQ4YM0W6lHU8haTYZPuC7du2qfWUdPivQQ4FVP5odgGGau3DhQpk5c6Z2ki/LsqRSpUqC95iuFXVhhoibkQAJxCPw5ptvStGiRbW9qBLPHT4kARIIAIGOHTv6lqyOBW8gEj+QeG3fvj0WTkruS4mXemFBPx/0rNCt0ap6JNWyCA2dM2bMqJZRUVqDJBaTP1HCU3w3VLTiWN2xY4d2kq+UKVNKly5dhP1+FD/IaB4JOEhg2LBhUqFCBa7g5SBTDhUYAnTUJwKQeFWrVk0riVcIlfGJnzfeeEMmTZqk3Y/gUIASu0elT5kyZVgWmxggn56HxAtNnf/880+fLOC0ThNAM+cHH3zQGHkUGj1jBYI77rhD8LfTvDievwSOHTsmw4cPl5MnT/prSBSzo9/P0qVLKfmKgh13IQHdCNStW9deLOHKK6/UzXTF7KU5JEACXhFAPzIdJV4hPkYnfg4dOiQDBw40TnJz9913C/qNQLYRCiTv/SewYsUKGTlypHHHm/9k/bMga9asonND58TI4QQblSGUfCVGSO/nIZ1Az4zTp09r50i6dOkEPfnY70e70NFgvwloND/65Y0aNUqyZ8+ukdU0lQRIIOgEOnXqpKXEKxQ3oxM/kHih5D3krCn3kHjxhE2taP7xxx8SFxen5VV2tUiqZQ1iev311xtZWVelShVp3LixIAmkFnVaEyuBkORr27ZtWla7opoV1bqsBIjuSOBeJKAyAVwxf/3118XU71aV2dM2EiCB6Ak888wzUrVqVS0lXiGvjU38vPbaazJt2jTjVlVC/5hSpUoZeSIaOih1vEcz5y1btshff/2lo/m0OQECDzzwgDRp0kTSpEmTwKv6P4WeKuj1g2XeLcvS3yF6cBmBEydOaLvKFySI5cuXl0WLFgmTP5eFlQ9IQGsC6DGHir7cuXPzd6zWkaTxJBAsAmhAj2qfa665RmvHjUz8HDhwQJAgMW0Vr3LlytmrnlDipdZ7Dj0pxo8fL+fOnVPLMFoTEwFIocxYYShxDLly5ZIBAwawsWbiiLR+BZ9NU6ZMER0lX0hMhpI/lH1pfRjSeBKwCaCvHD6TbrzxRiZ9bCL8hwRIQBcCSPrkzJlT+88uIxM/kHh9+eWXuhxLYdtJiVfYqDzb8MiRIwI5kBareHlGRf+J0EOrRIkSgpNP/b1J3APLsuSpp56S+++/33hfE6dg7iuQfGEhgE8++URLyRfef6HkDyt/zD1O6Zn5BO68805ZvHixsNLH/FjTQxIwjUDr1q0F7REyZMigvWvGJX7mz58vaGpp2qpKSGYVL17ckUyj9ketQg5A4oU+UpR4KRSUGE0pXLiwXVmHJrMxDqXF7qgg7NevH6t+tIhW5EaeOXPGlnzpuMoXvEXy57777rNlX9myZcNTvJEACWhE4K677hL07GLSR6Og0VQSIAGbACoVO3bsKLpLvOAMbkYlfvbt22dLvEyT3OBqPErMcIKGoPGmBgH0n5g6dSolXmqEwzErUMGVMWNGx8bTYSBcje3QoYOx/Yx0iIGbNr777rsyceJELSVf4ILkDyp/FixYIJAnWhZ7UoELbySgOoF77rlHFi5cyEbOqgeK9gWJAH2NgEDnzp21XsUrvqtGJX5QFbN79+74Pmr/mBIv9UL4008/CXrAUOKlXmxisQjNnKtVqyZBS7KimW7v3r0lX758gr9jYch91SMAyRequjZu3Cjnz59Xz8AwLELyB5U/L7/8stx88808TsNgxk1IwE8CSNaiCp+rd/kZhaTm5mskQAJJEWjbtq1UrlxZTJB4hfw0JvEzd+5cmTlzpnGreA0bNkzQSdyyeIUzdNCqcA+J165du4QSLxWi4YwNV111lZ3MM72hc2K04Dea4qdNmzaxTfi8xgRQCTtixAjROVkdSv689NJLUqhQISZ/ND4eabpCBFwwpWLFijJnzhxW6LnAlkOSAAm4T6BYsWKCSvjs2bO7P5mHMxiR+Pn++++NlHghy9i+ffvAVR94ePxHNdWrr74qL774orZXzqNyOgA7DRw4MPA/UmvVqiWPPvooP3MMPd5XrFgh48aN01byhbAg+QP5CC70oO8dHuN53mInwBFIwAkCaII6e/Zso+QRTnDhGCRAAvoQME3iFSJvROJn0KBB8s0334R8MuIecgtKvNQL5XfffSdYJUfH5ZHVo6mORUiyNm7cOPA9bnASDUlQUBpbq3MEemMJJF9Yvn/dunVJJa69MSaGWfD9iFX3UFGARGWaNGliGI27kgAJOEWgSZMmwqSPUzQ5DgmQgB8EIPGqVKmSoBLej/ndnFP7xA/0/viSMW0Vr+HDhwtWF7IsSrzcfANEOjYkXt9//z0lXpGCU3x7SJxM/ICPBnvevHkFyYF06dJFszv3UZwAJF+4qKCz5AuIkfy59dZbBb8BevToYeQPNPjJGwnoQAALIuAi7KhRo+zVbyyLv111iBttJAESuJwAKolNlHiFvNQ68bN3714jJV5Vq1aVZ599VnD1PRQo3vtPACcYuMKsa3PUqAkaviNOGqHl5fvtn0BblmV//pQuXZqfQf8gMe7fDz74QHCCpnvyx7IsQU+qXr16yfTp0yVPnjzs+2Pc0UqHVCdQpEgROwGLk6Wrr75aLItJH9VjRvtIgAQSJtClSxejZaopEnZbj2dxdQHSGz2sDc9KlKyPHDlSSclJeB6YudXXX38tOLnQ/UTJzOhE79Vtt90m3bp1E0qbLmeIVc3i4uLI5XIsxjyC5AvVi2vWrNFa8hUKSOrUqaVevXqycOFCQX8RPA69xnsSIAH3CNSuXVuwclfNmjUFVT/uzcSRSYAESMBdAii6qFChgtEVxCmiQKjELljVA9UXJkq8ChQowCsmShxl/xoxePBg+emnn/59gn8ZQQANnfljNeFQlitXTpo3b84kdMJ4tH82JPk6efKk9r7AASQrUaKN5E+7du2YtAQU3kjARQK4Mv7CCy8IfrMy2eoiaA5NAmoQMNoK9A3EgkqmreIVP2haJn5QfYGeHPjhGt8hnR/XqFFDnnnmGcorFAsiJAQLFiwQXCVXzDSaEwOBRo0a2dUBOGGMYRhjd0UPlb59+wZ+pTNjA/y3Y2vXrpURI0aIKZWMOGYzZcokqGbC6mX4AWdZlJ38HWr+TwKOEbjpppvslU179+4t2bJl44VKx8jqMhDtJAHzCKD6P2fOnMZ/nmmZ+IHEa9++fUYddZCaQOKFfgVGOaa5M1988YX07NlTuIqX5oGMZ36WLFns/mBs6BwPTLyHIU74fIr3Eh8aQAAVs0j8rFq1Sky6kALJ9FNPPWVLv+69915hcteAg5UuKEGgYsWKMm/ePHnyySflyiuvVMImGkECJEACsRBApU/58uWNlniF+GiX+JkxY4a8+uqrgh+sISdMuMcqK7fccosJrhjlAyRev/76q1E+0RmxV60KQmY/1lhbliX169eXatWq8eQ5VpiK7o+ED75/Tpw4oaiF0ZmFZA/kiqjWRBKIF1Wi4xjOXtwmGARatmwpc+fOFTT+5/spGDGnlyRgOoGSJUvaC5qgetF0X+FfCvyjy+2rr76yr9Ljh6ouNodj58MPP2z30uCqQuHQ8m6biRMnyuLFi41LMnpHUM2ZKlWqJE2aNGHvmjDDgxPoODZ6DpOWnputX79ehg0bFqvkSznnIf269tprZcKECTJ69GjJly8fV/1SLko0SHUCefPmlTFjxtifEddccw3fQ6oHjPaRAAmETQASr1y5chkv8QoB0Srxg0asP/74Y8h2I+7RjwDL6vLqiVrh3Lp1KyVeaoXEMWsGDBjA1UcipBla/Sy8z6kIB+fmvhNABS0SI8uXLzdK8gWwlmXZSd4WLVrIO++8Y1/Zo3QRZHgjgaQJ4PMefSffffddad26tUD6a1nsmZU0Nb5KAiSgC4EOHTrIfffdF6jFIFLoEpxp06bJa6+9Jn/99ZcuJodlJ/or3HTTTWFty428IwCJl2nSB+/oXZxJwbvu3btLsWLFeMUywtigcgJXRW6//Xayi5CdLpujkhbfR8eOHdPF5IjsROUa5NTwcdGiRXZjd65EFBFCbhwgAjVr1pS33nrLbv6O9w36ZgXIfbpKAiRgOIE777zTvhAUFIlXKJxaJH7QYLd///7GXYl85JFHBL0HTJd4hQ42Xe5R0owrw7gKrovNtDN5AgULFhQkfni1P3lWCW2Bk2RKvhIiY85zmzZtkuHDh4upSW/LsuzqnwceeMCW8WJBhTx58jCZac4hTE9iJIALI7jQOnv2bEEjZ1SlWxarfGLEyt1JgAQUI4CLmUHs9ZnCwzhEPRVW8Tp48GDU+6u4I0pmUVqPUloV7QuqTRs3bpS+ffvKmTNngorAWL8hFc2YMaOx/nnhWNWqVeWJJ54QJIG8mI9zeEsAye5x48YJpB2oAPJ2du9mw8UWJIAhY0GSH/d47J0FnIkE1CKQPXt26dGjh7z33o1LwS0AABAASURBVHv2BcmrrrpK8D5Ry0paQwIkoCEB5UyGxAuLPwTxez+FctGIZ9DkyZMFZdkmSrxuuOGGeN7yod8EkGQ8efKk32ZwfocJNGzYUJC0gNzD4aEDNRxOBFB9icR1oBwPkLNI+GCVryNHjhgnrY4fRiQw8+fPL+izt3DhQrvCAc/F346PScBkAk8++aQsXbpU+vTpI2jezO9Jk6Ptp2+cmwT8J1CqVKlASrzk4n8pLt4reff5558buYpXvXr1pFGjRiwvV+yog8Rh5cqVXMVLsbjEak7mzJkFDZ0zZMgQ61Dc/28COXLkEEq+/gZh8P9btmyxJV+nTp0y2Mt/XEP/KlTePvjgg/bJL74Hrr/+en4//4OH/xpMAE1N58+fL7jAin4X6dOn98ZbzkICJEACPhFAy4cgSrxCuJVO/ECa8csvv4RsNeIeTaQo8VIvlB9++KHgeDt79qx6xtGimAigQiXIH/IxwUtgZ8uypHnz5oIyWVQAJbAJn9KcACRfkyZNkrffftu43nqJhQbHMk5827ZtK8uWLbNXMbryyisT25zPO0CAQ/hDoGjRonbT5jfeeEMeffRRwXGOBKg/1nBWEiABEvCGQMeOHaVs2bKBWsUrPlllEz8TJkyQJUuWGFdqjhVFrrvuuvhx4GMfCUDagKRPEK5u+4jZl6krVKhg9yvgiiTO4occBlVUQdRHO0tS3dHwuYjql0OHDrn9PawUBBzbt912m6Dx89q1a6Vr166CKjfLYoNbpQJFYyImcO+99woSuqtWrRL0uMiaNatQ1hUxRu5AAiSgIYHSpUsLLuygAEND8x0zWcnEz7Zt22wpAX54OuapAgOhKSpuvLKiQDAuMWHYsGGCih9c5b7kaf5pAAH0bGJDZ3cCCWkAvkSdS6q5YydHjZ7AZ599FhjJ16WU8B0N+VeRIkXsSlAkgPr37y/58uWjBOxSUPxbCwLVq1eXOXPmyFtvvWVXayLhgwSnFsbTSBIgARJwgEDQJV4hhEomflB98dtvv4VsNOIeDfMo8VIvlO+//759YkOJl0KxcciULl26SLFixXii5hDP+MPg5BgnwzfffDMZx4djyGMkw6dNmyZvvvmmBPEzEsc4Epu33HKLdOvWTdasWWNXAuFzBdIwQ8JMNwwkgMRl/fr1bbnmvHnzBL0ls2TJwhUZDYw1XSIBEkiaQKdOnQIv8QoRUi7xM3bsWLvBommreGGVFCyXGQLP++QJuL0FVu9CkhH3bs/F8b0lUKBAAenZs2egdbxeEMfJBZI/uPdiPs7hPQFU3kKifODAgUBJvi4lbVmWIAEEyReq3FavXm03xEWvAEplLiXFv/0mABlDq1atBBe1XnjhBXs1S/Tw4XHqd2Q4PwmQgB8EIPFq06aNoNLRj/lVmzOFSgZhJZG4uDg5f/78pWZp/3fDhg3l8ccf51VxxSI5ZMgQ+fjjjwN7MqNYOBw1BxKvTJkyOTomB0uYwCOPPCK1atVir4iE8Rjx7BdffGE3gw16ktyyLLtiAisFPvXUU7JixQqZNWuWfXJN6YwRh7q2TqDyElVpSEqiuvzuu+8WfAeyMk3bkNJwEgg6AUf879Gjh+TKlUssy3JkPN0HUSrxg+qL33//XXeml9mP1YTGjBljXy287AU+8JUAVm1BXHA121dDOLnjBJ588kmpVq0aExGOk014QFxJxmc3VkRKeAs+qzsBSL5mzJghWAUoiJKvhOKH4x7HPCQ0kMItWLBAGjVqJJB1WxZ/YCbEjM85TwCrK6JPYagPVaFChQTHJWSKzs/GEUnADwKckwSiI/Dcc8/JPffcI6xK/5efMomfUaNGCU7GTZR4XX311f8S51++E0D/KKxIxFW8fA+F4wagpH3w4MH2D1/HB+eAiRLA1eY+ffrwyzVRQvq/gCQ5JF8//PCDIBGkv0fOeIAEEFa3e+ihhwTSmk8//VReeuklQX8VfPdbFpNAzpDmKCECSPYg2b5p0ya7hw/6V+CKNk5ujE74hADwngRIgASSIVCmTBlp3bo1JV7xOCmR+MGXl4kSr6ZNm0qdOnVYXhbvoPP7ISReW7dupcTL70C4MD8SeujDYVk82XIBb6JDWpYlWB64ZMmSQmlBopi0f+HLL7+0m+Ezaf7fUOK4x4k3TsBRdThz5kxBEmj69OlSt25dgTzMsvi59F9ykT8TxD3iJ3sg68Ln7RVXXGHLDy2Lx1YQjwv6TAIkkDABrOKF72PL4mfjpYSUSPzgysXx48cvtUv7v3Pnzi2QEqEhpPbOGOQASvInT54suHptkFt05W8C999/vzz99NOUVf7Nwo//0eMEiTdUP/gxP+d0nwAqfdDTZv78+Sqs8uW+w1HMYFmWLTPF++DGG2+Uxo0by8svvyzoYThlyhSpXbu2ZMyYMYqRuUvQCCSX7GF1T9COCPpLAiQQDoHOnTsL+pzhYkw42wdpG98TPygdX758uXHMsYoXZCfGOaaxQz///LMgycir1RoHMQnT0dA5Q4YMSWzBl9wmcN9990nTpk09Tr657RXHv5QAFl/A99vevXsp+boUTAJ/W9a/SaC8efMKGkLPmzfPrgSaOHGiQB7GH6YJgAvwU0j24LsMlfBvv/22xK/sYbInwAcHXScBEkiWABI+lHgljsnXxM9HH30kcQau4tW8eXP7qp5lsbws8UPP+1fQ+2XHjh2UeHmP3vUZu3TpIsWLF+fKea6TTnoCyF2QXL322mspcU0aldav7t6925Z8BX2Vr0iCaFmWLclBJdCtt94qLVq0EDSEhhwMvYHwu+GOO+6wq4UiGZfb6k2gQIEC0qBBAxk7dqxs3rzZ7tnTtWtXuVTGxWSP3jGm9SRAAt4RgMQLCytZFs/BE6Lua+IHJwimVV+gySkkXpA9JAScz7lPIKEZFi5cKOi1gKvVCb3O5/QlkD9/funVq5fghEpfL8yxPEuWLNK/f3/Gw5yQ/scTLMIwZ84cQfXKmTNn/vM6n0iagGX9mwQqWLCgXSU3YcIE2bBhg10N9OKLL0qrVq2kcOHC7JmVNErtXkW8UfkF2d8nn3wiqOzBinm4Ql2iRAkJ9exhske70NJgEiABnwngIjCaOrOSNvFA+Jb4GTp0qNx5553Ss2dPt26+jIvVySg3SfyA8+OVffv2yddffy0dO3b05Zgw7RhXzZ9hw4axZ4Yfb6wk5mzYsKH06NGD7zfDvt8ufe+jKuHgwYPCVb6SeCOE8RJO8LEyGH6o4rcDkj2NGjWyK0CQFPj8889l9uzZ0rJlS0HSAFV1YQzLTRQhgJg1a9ZM0OwbFceIKfoM4rnixYsLWgIg9rhYiGNBEbNphgYE2rRpw+9Yg79jL/2+5d/h5Qpwnpc1a1YN3r3+mehb4ue5556T3r17S79+/Yy61axZkxIH/47nBGe+/vrrBVlg0441+vPPZ0eNGjUoj0jwyPfvSZzIotyWx+g/x6ipHBBj9K7hCatz7zWwxPsHyQAkgiAFeuKJJ2TcuHF2g+gvvvhC0GAbVSO33HILK4KcQ+/ISKFED2K0a9cuW76Ffk5IhhcqVMiu6EFsEWPE2pFJOUjgCKBCrG/fvkadPzn/PWn29y95/Te+1113Hc/Bk/k09C3xg9WuTLzxalwyR5wPL1uWZTebNfF4o09pmPTx4T0VzpS4gs3jM43xnz08eQ3n3RD9NuCLJAGSBUgEQdr65JNPCqpGtm/fLt9++628//779uP27dtLtWrVBJJz/haJnnk4e2bPnl3uvfdeuxIL/XlWrFgh+/fvF8i3kOhBjPLlyyeZMmUSxA4xRCzDGZvbOEzAwOH4/Wr+dyt/P0UeY8tiX5/kPu5SJLcBXycBEiABEiABEiABEvCfgGVZdrIbyQQkglDRWr58eYF0aPjw4bJo0SKBpOjQoUOyceNGeyl5VFfXqVNHihQpYich/PfCHwuimTVz5sxSqlQpadKkid3QfOnSpbJnzx755ptvZPny5XYlFqovKlSoIGgomjFjRpsxEj2WxZOQaJhzHxIgARIgAXcIMPHjDleOSgIkQAIkQAIkoB4BoyyyLMuWe6ECAMkgNLlHQihr1qx2H8X69etLnz595JVXXpGPP/5YDh8+LJCLvfHGG4Jei88884w8/PDD9ra5cuUK5MqI4HXjjTfaDMALzelfe+012blzp93Das2aNTJ16lS7T2CoogpNmNOnT28necAeFVaWxUSP8D8SIAESIAFlCTDxo2xoaBgJkAAJkIB7BDgyCZhNANIiVJ5AMhBKCEF6hD40SPag1+KYMWNk/vz5snbtWkFPmqNHj9qLIeDxq6++KqNHj5bOnTsL+gyVL19eIDVD0sOy1E5yoGky+k9hhRf42rx5c7sRLmRZc+fOlZUrVwqkcj/99JOgOur//u//bAZowozeVbVq1bKbaWMcJnjMfp/QOxIgARIICgEmfoISafpJAiSQMAE+SwIkQAIBIoCEECpUUKkSqhJCcgMyJSR1kDApW7asQB7Wtm1bGTx4sLz44ouybNky2bp1qyBZgn42W7ZssZ+DvAxJolmzZsm0adNk/Pjx8vzzz8ugQYPsaiOsAIf+Q61atZKmTZsKqmoeffRRQWP+SpUqSbly5aRixYry4IMPSu3ateXxxx+Xxo0bS4sWLeTZZ5+1E09Y1WbAgAG23ArJqilTptgrZc2ZM0dQnfPWW2/JmjVr7Cqdn3/+WX788UfBimirVq2yE1sTJkywG+FClvXYY4/J/fffL1hB7dprr7X78KDqBwyQIEOiDAkzy1I7uRWgQ5aukgAJkAAJOECAiR8HIHIIMwjQCxIgARIgARIIOgHLsmzJF5IfSA4hEYKECBIjSJAgQZQjRw4pVqyYPPDAA/LQQw/ZSSI0NMZqY0jwINGDhA/6Cw0cONBO2IwdO1YSSti89957gt45kJ/NmzdPXnrppcsSSEg8YQWbbt262XKrNm3ayNNPPy1YKatevXqC6hwkjZCsQjUTGi+jsgm2wmbYjgQX/IA/8AuJLyTALMsS/kcCJEACJEACQSDAxM9/o8xnSIAESIAESIAESIAEEiFgWZadHEICBTckU0I3JFdCNyRbcEPiJXRDIiZ0Q2ImdAs9h/vQtrjH/qFbaFzch+bDPWzAzbKYyBH+RwIkQAIkECmBQGzPxE8gwkwnSYAESIAESIAESIAESIAESIAEEifAV0jAXAJM/JgbW3pGAiRAAiRAAiRAAiRAAiQQKQFuTwIkQAKGEWDix7CA0h0SIAESIAESIAESIAFnCHAUEiABEiABEjCBABM/JkSRPpAACZAACZAACbhJgGOTAAmQAAmQAAmQgLYEmPjRNnQ0nARIgARIwHsCnJEESIAESIAESIAESIAE9CLAxI9e8aK1JEACqhCgHSRAAiRAAiRAAiRAAiRAAiSgAQEmfjQIEk1UmwCtIwESIAESIAESIAESIAESIAESIAFVCTDx41xkOBIJkAAJkAAJkAAJkAAJkAAJkADEv6//AAAAsElEQVQJkID5BLTykIkfrcJFY0mABEiABEiABEiABEiABEiABNQhQEtIQH0CTPyoHyNaSAIkQAIkQAIkQAIkQAIkoDoB2kcCJEACihJg4kfRwNAsEiABEiABEiABEiABPQnQahIgARIgARJQiQATPypFg7aQAAmQAAmQAAmYRIC+kAAJkAAJkAAJkIDvBJj48T0ENIAESIAESMB8AvSQBEiABEiABEiABEiABPwh8P8AAAD//3Z3F/4AAAAGSURBVAMAVOjcBH8cPIkAAAAASUVORK5CYII=" style="width:100px;height:auto;display:block;margin:0 auto;" alt="XUL"/>
    </div>

    <div class="sb-section">
      <div class="sb-label">Por categoría</div>
      <div class="sb-item" onclick="setCpv('')">
        <span class="sb-dot blue"></span> Todos los CPVs
      </div>
      <div class="sb-item" onclick="setCpv('79341000')">
        <span class="sb-dot blue"></span> Publicidad
      </div>
      <div class="sb-item" onclick="setCpv('79340000')">
        <span class="sb-dot blue"></span> Marketing
      </div>
      <div class="sb-item" onclick="setCpv('79952000')">
        <span class="sb-dot green"></span> Eventos
      </div>
      <div class="sb-item" onclick="setCpv('79956000')">
        <span class="sb-dot green"></span> Ferias
      </div>
      <div class="sb-item" onclick="setCpv('79822500')">
        <span class="sb-dot amber"></span> Diseño gráfico
      </div>
      <div class="sb-item" onclick="setCpv('72413000')">
        <span class="sb-dot amber"></span> Web
      </div>
      <div class="sb-item" onclick="setCpv('79416000')">
        <span class="sb-dot red"></span> RRPP
      </div>
      <div class="sb-item" onclick="setCpv('79800000')">
        <span class="sb-dot" style="background:#9CA3AF"></span> Impresión
      </div>
    </div>

  </div>

  <!-- Main -->
  <div class="main">

    <!-- Topbar -->
    <div class="topbar">
      <div class="topbar-search">
        <span class="search-icon">🔍</span>
        <input type="text" id="search" placeholder="Buscar licitación u órgano…" oninput="buscar()">
      </div>
      <div class="topbar-sep"></div>
      <div class="topbar-actions">
        <button class="btn-primary" id="btn-run" onclick="ejecutarAgente()">
          <span class="btn-label">▶ Buscar ahora</span>
          <div class="spinner"></div>
        </button>
      </div>
    </div>

    <!-- Page head -->
    <div class="page-head">
      <div class="page-title-row">
        <div class="page-title" id="page-title">Licitaciones activas</div>
        <div class="page-badge">IMAGINE COMUNICACIÓN ANDALUZA</div>
      </div>
    </div>

    <!-- Stats -->
    <div class="stats-row">
      <div class="stat-card blue">
        <div class="s-label">📋 Licitaciones</div>
        <div class="s-val" id="stat-total">—</div>
        <div class="s-sub" id="stat-update">Última actualización: —</div>
      </div>
      <div class="stat-card green">
        <div class="s-label">💰 Importe total</div>
        <div class="s-val" id="stat-importe">—</div>
        <div class="s-sub">mercado estimado</div>
      </div>
      <div class="stat-card amber">
        <div class="s-label">📊 Importe medio</div>
        <div class="s-val" id="stat-medio">—</div>
        <div class="s-sub">por licitación</div>
      </div>
      <div class="stat-card red">
        <div class="s-label">🏆 Mayor contrato</div>
        <div class="s-val" id="stat-max">—</div>
        <div class="s-sub">valor más alto</div>
      </div>
    </div>

    <!-- Toolbar -->
    <div class="toolbar">
      <select class="tb-select" id="filter-cpv" onchange="cargarTodo()">
        <option value="">Todos los CPVs</option>
        <option value="79340000">Publicidad y marketing</option>
        <option value="79341000">Publicidad</option>
        <option value="79342000">Marketing</option>
        <option value="79341200">Gestión publicidad</option>
        <option value="79341400">Campañas</option>
        <option value="79342200">Promoción</option>
        <option value="79341100">Colocación publicidad</option>
        <option value="92210000">Radio</option>
        <option value="92220000">Televisión</option>
        <option value="79822500">Diseño gráfico</option>
        <option value="74812240">Vídeos publicidad</option>
        <option value="92111250">Vídeos información</option>
        <option value="79800000">Impresión</option>
        <option value="72413000">Diseño web</option>
        <option value="72415000">Alojamiento web</option>
        <option value="79413000">Consultoría marketing</option>
        <option value="79416000">Relaciones públicas</option>
        <option value="79416200">Asesoramiento RRPP</option>
        <option value="79952000">Organización eventos</option>
        <option value="79956000">Ferias y exposiciones</option>
      </select>
      <select class="tb-select" id="filter-min" onchange="cargarTodo()">
        <option value="0">Cualquier importe</option>
        <option value="10000">+10.000 €</option>
        <option value="50000">+50.000 €</option>
        <option value="100000">+100.000 €</option>
        <option value="500000">+500.000 €</option>
      </select>
      <select class="tb-select" id="filter-plazo" onchange="renderGrid(allLics)">
        <option value="">Cualquier plazo</option>
        <option value="3">Vence en 3 días</option>
        <option value="7">Vence en 7 días</option>
        <option value="15">Vence en 15 días</option>
        <option value="30">Vence en 30 días</option>
      </select>
      <div class="count-pill" id="count-pill">0 resultados</div>
    </div>

    <!-- Cards -->
    <div class="cards-area">
      <div class="empty-state" id="empty-state"><div class="e-icon">⏳</div><p>Cargando licitaciones…</p></div>
      <div class="card-grid" id="card-grid" style="display:none;"></div>
    </div>

  </div><!-- /main -->
</div><!-- /app -->

<!-- Modal detalle -->
<div class="modal-bg" id="modal-det" onclick="if(event.target===this)this.classList.remove('open')">
  <div class="modal">
    <div class="modal-head">
      <div class="modal-head-icon">📋</div>
      <h2 id="m-titulo">—</h2>
      <button class="modal-close" onclick="document.getElementById('modal-det').classList.remove('open')">✕</button>
    </div>
    <div class="modal-body">
      <div class="mf"><label>Órgano contratante</label><p id="m-organo">—</p></div>
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px;">
        <div class="mf"><label>Importe estimado</label><p class="mf-val green" id="m-importe" style="color:#059669">—</p></div>
        <div class="mf"><label>Plazo presentación</label><p class="mf-val" id="m-plazo">—</p></div>
      </div>
      <div class="mf"><label>CPVs</label><p id="m-cpvs" style="white-space:pre-line;font-size:12px;color:#555;">—</p></div>
      <div class="mf"><label>Detectada el</label><p id="m-fecha">—</p></div>
    </div>
    <div class="modal-footer">
      <a id="m-enlace" class="btn-link" target="_blank">🔗 Ver en PLACSP</a>
    </div>
  </div>
</div>

<!-- Modal análisis Grok -->
<div class="modal-bg" id="modal-analisis" onclick="if(event.target===this)this.classList.remove('open')">
  <div class="modal modal-analisis">
    <div class="modal-head">
      <div class="modal-head-icon" style="background:#F0F0F0;color:#111;font-size:16px;">✦</div>
      <div style="flex:1;min-width:0;">
        <div style="font-size:10px;font-weight:600;color:#999;text-transform:uppercase;letter-spacing:.6px;margin-bottom:2px;">Análisis Grok</div>
        <h2 id="ma-titulo" style="font-size:13px;color:#555;font-weight:400;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">—</h2>
      </div>
      <button class="modal-close" onclick="document.getElementById('modal-analisis').classList.remove('open')">✕</button>
    </div>
    <div class="modal-body" id="ma-body">
      <div class="analysis-loading"><div class="big-spin"></div><p>Analizando con Grok…</p></div>
    </div>
    <div class="modal-footer" style="display:flex;gap:8px;align-items:center;" id="ma-footer">
      <a id="ma-enlace" class="btn-link" target="_blank" style="display:none;">🔗 Ver en PLACSP</a>
      <button id="btn-docx" onclick="descargarDocx()" style="display:none;background:#1A56DB;color:#fff;border:none;border-radius:8px;padding:8px 16px;font-size:13px;font-weight:500;cursor:pointer;display:none;align-items:center;gap:6px;">📄 Descargar DOCX</button>
      <button onclick="document.getElementById('modal-analisis').classList.remove('open')" style="margin-left:auto;background:#fff;border:1px solid #EBEBEB;border-radius:8px;padding:8px 16px;font-size:13px;cursor:pointer;">Cerrar</button>
    </div>
  </div>
</div>

<!-- Modal log -->
<div class="modal-bg" id="modal-log" onclick="if(event.target===this)this.classList.remove('open')">
  <div class="modal" style="max-width:720px">
    <div class="modal-head">
      <div class="modal-head-icon">📄</div>
      <h2>Resultado última ejecución</h2>
      <button class="modal-close" onclick="document.getElementById('modal-log').classList.remove('open')">✕</button>
    </div>
    <div class="modal-body">
      <pre class="log-pre" id="log-content">Sin datos aún. Pulsa "Buscar ahora" para ejecutar el agente.</pre>
    </div>
  </div>
</div>

<!-- Modal config -->
<div class="modal-bg" id="modal-config" onclick="if(event.target===this)this.classList.remove('open')">
  <div class="modal" style="max-width:400px">
    <div class="modal-head">
      <div class="modal-head-icon">⚙️</div>
      <h2>Configuración</h2>
      <button class="modal-close" onclick="document.getElementById('modal-config').classList.remove('open')">✕</button>
    </div>
    <div class="modal-body">
      <div class="mf"><label>Estado del agente</label><p id="cfg-estado">—</p></div>
      <div class="mf"><label>Licitaciones en caché</label><p id="cfg-cache">—</p></div>
      <div class="mf" style="margin-top:16px;padding-top:16px;border-top:1px solid #F0F0F0;">
        <label>Filtros activos</label>
        <p style="font-size:12px;color:#555;margin-top:4px;">Solo activas y relevantes para IMAGINE. Puedes desactivarlos usando los parámetros de la URL: <code style="background:#F5F5F5;padding:2px 5px;border-radius:4px;">/api/licitaciones?activas=0&relevantes=0</code></p>
      </div>
    </div>
  </div>
</div>

<!-- Toast -->
<div class="toast" id="toast"></div>

<script>
const CPV_NAMES = {
  "79340000":"Publicidad y marketing","79341000":"Publicidad","79342000":"Marketing",
  "79341200":"Gestión publicidad","79341400":"Campañas","79342200":"Promoción",
  "79341100":"Colocación publicidad","92210000":"Radio","92220000":"Televisión",
  "79822500":"Diseño gráfico","74812240":"Vídeos publicidad","92111250":"Vídeos info",
  "79800000":"Impresión","72413000":"Diseño web","72415000":"Alojamiento web",
  "79413000":"Consultoría marketing","79416000":"Relaciones públicas","79416200":"Asesoramiento RRPP",
  "79952000":"Organización eventos","79956000":"Ferias y exposiciones",
  "79950000":"Eventos (gen.)","79955000":"Desfiles","92111200":"Producción audiovisual",
  "92111210":"Producción vídeo","92111220":"Audiovisual"
};

let allLics = [];
let activeTab = 'todas';

function fmt(n) {
  if (!n || n == 0) return 'N/D';
  return new Intl.NumberFormat('es-ES',{style:'currency',currency:'EUR',maximumFractionDigits:0}).format(n);
}

function fmtDate(d) {
  if (!d || d==='N/D') return 'N/D';
  try { return new Date(d).toLocaleDateString('es-ES',{day:'2-digit',month:'short',year:'numeric'}); }
  catch { return d; }
}

function plazoDays(d) {
  if (!d || d==='N/D') return null;
  try { return Math.ceil((new Date(d) - new Date()) / 86400000); }
  catch { return null; }
}

function dueChip(d) {
  const days = plazoDays(d);
  if (days === null) return '';
  if (days < 0)   return `<span class="due-chip lejano">Cerrado</span>`;
  if (days <= 7)  return `<span class="due-chip urgente">⚡ Vence en ${days}d</span>`;
  if (days <= 14) return `<span class="due-chip pronto">⏳ Vence en ${days}d</span>`;
  if (days <= 30) return `<span class="due-chip ok">📅 ${fmtDate(d)}</span>`;
  return `<span class="due-chip lejano">📅 ${fmtDate(d)}</span>`;
}

function renderCard(l) {
  const imp = parseFloat(l.importe) || 0;
  const cpvs = [...new Set((l.cpvs||[]).map(c=>c.slice(0,8)))]
    .filter(c => CPV_NAMES[c]).slice(0,2)
    .map(c => `<span class="cpv-chip">${CPV_NAMES[c]}</span>`).join('');
  const impHtml = imp > 0
    ? `<span class="lc-importe">${fmt(imp)}</span>`
    : `<span class="lc-importe nd">N/D</span>`;

  const pubDate  = l.fecha_deteccion ? `📅 Publicada: ${fmtDate(l.fecha_deteccion)}` : '';
  const plazoTxt = l.plazo ? `⏱️ Plazo: ${fmtDate(l.plazo)}` : '';

  const ld = JSON.stringify(l).replace(/"/g,'&quot;');
  return `<div class="lic-card">
    <div class="lc-top" onclick="verDetalle(${ld})" style="cursor:pointer;">
      <div class="lc-due">${dueChip(l.plazo)}<span class="lc-menu">···</span></div>
      <div class="lc-title">${l.titulo||'Sin título'}</div>
      <div class="lc-organo">${l.organo||'Órgano desconocido'}</div>
      <div class="lc-dates">${[pubDate, plazoTxt].filter(Boolean).join('<span class="lc-date-sep">·</span>')}</div>
    </div>
    <div class="lc-bottom">
      <div class="lc-cpvs">${cpvs||'<span class="cpv-chip" style="background:#F5F5F5;color:#999;">Sin CPV</span>'}</div>
      ${impHtml}
    </div>
    <div class="lc-actions">
      <button class="btn-analizar" id="btn-a-${l._id?.slice(-8)||Math.random().toString(36).slice(2)}" onclick="analizarLic(${ld}, this)">
        <span class="a-label">✦ Analizar con Grok</span>
        <div class="a-spin"></div>
      </button>
      ${l.enlace ? `<button class="btn-ver" onclick="window.open('${l.enlace}','_blank')">🔗 PLACSP</button>` : ''}
    </div>
  </div>`;
}

function applyTab(lics) {
  const maxDays = parseInt(document.getElementById('filter-plazo')?.value || '0');
  if (maxDays > 0) return lics.filter(l => { const d=plazoDays(l.plazo); return d!==null && d>=0 && d<=maxDays; });
  return lics;
}

function renderGrid(lics) {
  const grid = document.getElementById('card-grid');
  const emptyEl = document.getElementById('empty-state');
  const shown = applyTab(lics);
  document.getElementById('count-pill').innerHTML = `<strong>${shown.length}</strong> resultado${shown.length!==1?'s':''}`;
  if (shown.length === 0) {
    grid.style.display = 'none';
    emptyEl.style.display = 'flex';
    emptyEl.innerHTML = '<div class="e-icon">🔍</div><p>No hay licitaciones con ese filtro</p>';
    return;
  }
  emptyEl.style.display = 'none';
  grid.style.display = 'grid';
  grid.innerHTML = shown.map(renderCard).join('');
}


function setCpv(cpv) {
  document.getElementById('filter-cpv').value = cpv;
  cargarTodo();
}

function buscar() { cargarTodo(); }

async function cargarTodo() {
  const q   = document.getElementById('search').value;
  const cpv = document.getElementById('filter-cpv').value;
  const min = document.getElementById('filter-min').value;
  let params = `activas=1&relevantes=1`;
  if (q)   params += `&q=${encodeURIComponent(q)}`;
  if (cpv) params += `&cpv=${cpv}`;
  if (min && min!='0') params += `&min=${min}`;

  const r = await fetch(`/api/licitaciones?${params}`);
  const d = await r.json();
  allLics = d.licitaciones;

  const urgentes    = allLics.filter(l => { const d=plazoDays(l.plazo); return d!==null&&d>=0&&d<=7; });
  const semana      = allLics.filter(l => { const d=plazoDays(l.plazo); return d!==null&&d>=0&&d<=14; });
  const grandes     = allLics.filter(l => (parseFloat(l.importe)||0)>=100000);

  renderGrid(allLics);
  cargarStats();
}

async function cargarStats() {
  const r = await fetch('/api/stats');
  const d = await r.json();
  document.getElementById('stat-total').textContent   = d.total;
  document.getElementById('stat-importe').textContent = fmt(d.importe_total);
  document.getElementById('stat-medio').textContent   = fmt(d.importe_medio);
  document.getElementById('stat-max').textContent     = fmt(d.importe_max);
  document.getElementById('stat-update').textContent  = d.actualizado
    ? 'Últ: ' + new Date(d.actualizado).toLocaleString('es-ES',{day:'2-digit',month:'2-digit',hour:'2-digit',minute:'2-digit'})
    : '—';
}

function formatAnalysis(text) {
  const SECTION_RE = /^(🎯|🛠️|📋|⚡|💡|✅|🚀)\s/;
  const lines = text.split('\n');
  let html = '<div class="analysis-body">';
  let inSection = false, inUl = false;

  const closeUl  = () => { if (inUl)      { html += '</ul>'; inUl = false; } };
  const closeSec = () => { closeUl(); if (inSection) { html += '</div>'; inSection = false; } };

  for (let line of lines) {
    line = line.trim();
    if (!line) continue;

    if (SECTION_RE.test(line) || (/^\*\*/.test(line) && line.endsWith('**'))) {
      closeSec();
      const title = line.replace(/\*\*/g,'').trim();
      html += `<div class="analysis-section"><h3>${title}</h3>`;
      inSection = true;
    } else if (/^[-•*]\s/.test(line) || /^\d+\./.test(line)) {
      if (!inUl) { html += '<ul>'; inUl = true; }
      const content = line.replace(/^[-•*\d\.]\s*/,'').replace(/\*\*(.*?)\*\*/g,'<strong>$1</strong>');
      html += `<li>${content}</li>`;
    } else {
      closeUl();
      const content = line.replace(/\*\*(.*?)\*\*/g,'<strong>$1</strong>');
      html += `<p>${content}</p>`;
    }
  }
  closeSec();
  html += '</div>';
  return html;
}

let _currentLic = null;
let _currentAnalysis = null;

async function analizarLic(l, btn) {
  _currentLic = l;
  _currentAnalysis = null;
  document.getElementById('btn-docx').style.display = 'none';
  document.getElementById('ma-titulo').textContent = l.titulo || '—';
  const enlaceEl = document.getElementById('ma-enlace');
  enlaceEl.href = l.enlace || '#';
  enlaceEl.style.display = l.enlace ? 'inline-flex' : 'none';
  document.getElementById('ma-body').innerHTML =
    '<div class="analysis-loading"><div class="big-spin"></div><p>Analizando con Groq…</p></div>';
  document.getElementById('modal-analisis').classList.add('open');

  if (btn) { btn.disabled = true; btn.classList.add('loading'); }

  try {
    const r = await fetch('/api/analizar', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify({
        titulo: l.titulo, organo: l.organo,
        importe: l.importe, plazo: l.plazo,
        cpvs: l.cpvs, enlace: l.enlace
      })
    });
    const d = await r.json();
    if (d.ok) {
      _currentAnalysis = d.analysis;
      document.getElementById('ma-body').innerHTML = formatAnalysis(d.analysis);
      document.getElementById('btn-docx').style.display = 'inline-flex';
    } else {
      const isApiKey = d.msg && d.msg.includes('API key');
      document.getElementById('ma-body').innerHTML = isApiKey
        ? `<div class="no-api-warn"><strong>API key no configurada.</strong></div>`
        : `<p style="color:#DC2626;font-size:13px;">Error: ${d.msg}</p>`;
    }
  } catch(e) {
    document.getElementById('ma-body').innerHTML = `<p style="color:#DC2626;font-size:13px;">Error de conexión: ${e}</p>`;
  }
  if (btn) { btn.disabled = false; btn.classList.remove('loading'); }
}

async function descargarDocx() {
  if (!_currentAnalysis || !_currentLic) return;
  const btn = document.getElementById('btn-docx');
  btn.textContent = '⏳ Generando…';
  btn.disabled = true;
  try {
    const r = await fetch('/api/descargar-analisis', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify({
        titulo: _currentLic.titulo, organo: _currentLic.organo,
        importe: _currentLic.importe, plazo: _currentLic.plazo,
        analysis: _currentAnalysis
      })
    });
    const blob = await r.blob();
    const url  = URL.createObjectURL(blob);
    const a    = document.createElement('a');
    a.href = url;
    a.download = r.headers.get('Content-Disposition')?.match(/filename="?([^"]+)"?/)?.[1] || 'analisis.docx';
    a.click();
    URL.revokeObjectURL(url);
  } catch(e) { toast('❌ Error al descargar', 'err'); }
  btn.textContent = '📄 Descargar DOCX';
  btn.disabled = false;
}

function verDetalle(l) {
  document.getElementById('m-titulo').textContent  = l.titulo||'—';
  document.getElementById('m-organo').textContent  = l.organo||'—';
  document.getElementById('m-importe').textContent = fmt(parseFloat(l.importe)||0);
  document.getElementById('m-plazo').textContent   = fmtDate(l.plazo) + (plazoDays(l.plazo)!==null ? ` (${plazoDays(l.plazo)} días)` : '');
  document.getElementById('m-fecha').textContent   = l.fecha_deteccion||'—';
  const cpvLines = [...new Set((l.cpvs||[]).map(c=>c.slice(0,8)))]
    .map(c => `${c}  ${CPV_NAMES[c]||''}`.trim()).join('\n');
  document.getElementById('m-cpvs').textContent = cpvLines||'—';
  document.getElementById('m-enlace').href = l.enlace||'#';
  document.getElementById('modal-det').classList.add('open');
}

function toast(msg, type='') {
  const el = document.getElementById('toast');
  el.textContent = msg;
  el.className = 'toast show ' + type;
  setTimeout(()=>el.className='toast', 4000);
}

async function ejecutarAgente() {
  const btn = document.getElementById('btn-run');
  btn.disabled = true;
  btn.classList.add('loading');
  toast('⏳ Buscando licitaciones, puede tardar 1-2 min…');
  try {
    const r = await fetch('/api/ejecutar',{method:'POST'});
    const d = await r.json();
    if (!d.ok && d.msg) { toast('❌ ' + d.msg, 'err'); btn.disabled=false; btn.classList.remove('loading'); return; }
    // polling hasta que termine
    let intentos = 0;
    const poll = setInterval(async () => {
      intentos++;
      const s = await fetch('/api/estado').then(x=>x.json());
      if (!s.agente_corriendo) {
        clearInterval(poll);
        btn.disabled = false;
        btn.classList.remove('loading');
        const res = s.ultimo_resultado || {};
        if (res.stdout) document.getElementById('log-content').textContent = res.stdout;
        if (res.ok) { toast('✅ Búsqueda completada', 'ok'); await cargarTodo(); }
        else { toast('❌ ' + (res.msg || res.stderr || 'Error al ejecutar'), 'err'); }
      } else if (intentos > 60) {
        clearInterval(poll);
        btn.disabled = false;
        btn.classList.remove('loading');
        toast('⚠️ Tiempo de espera agotado', 'err');
      }
    }, 3000);
  } catch(e) { toast('❌ Error de conexión','err'); btn.disabled=false; btn.classList.remove('loading'); }
}

async function cargarEstado() {
  const r = await fetch('/api/estado');
  const d = await r.json();
  document.getElementById('cfg-estado').textContent   = d.agente_corriendo ? 'En ejecución…' : 'Inactivo';
  document.getElementById('cfg-cache').textContent    = d.licitaciones_en_cache + ' IDs guardados';
}

cargarTodo();
cargarEstado();
</script>
</body>
</html>"""

@app.get("/")
def index():
    return render_template_string(HTML)


if __name__ == "__main__":
    print("\n  Agente Licitador — IMAGINE COMUNICACIÓN ANDALUZA S.L.U")
    print("  → http://localhost:5000\n")
    app.run(debug=False, port=5000)
